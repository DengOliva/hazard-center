import io
import json
import math
import os
import re
import sqlite3
import uuid
from io import BytesIO
from datetime import date, datetime, timedelta
from pathlib import Path

from flask import Flask, jsonify, request, send_file, send_from_directory
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

from data_admin import bp as admin_bp
from meeting import bp as meeting_bp
from collection import bp as collection_bp

ROOT = Path(__file__).resolve().parent
DATA_DIR = Path(os.environ.get("DATA_DIR", ROOT / "data"))
DB_PATH = DATA_DIR / "hazards.db"
UPLOAD_DIR = DATA_DIR / "uploads"
TRAINING_LEDGER_DIR = DATA_DIR / "training_ledger"
BRAKE_LEDGER_DIR = DATA_DIR / "brake_ledger"
FEEDBACK_DIR = DATA_DIR / "feedback"
SEED_DIR = ROOT / "seed"
TRAINING_SCHEDULE_FILE = SEED_DIR / "2026年7月安全培训安排表.xlsx"
TRAINING_OVERRIDES_FILE = DATA_DIR / "training_overrides.json"
TRAINING_MATERIALS_FILE = DATA_DIR / "training_materials.json"
TRAINING_EDIT_PASSWORD = os.environ.get("TRAINING_EDIT_PASSWORD", "@q")
MATERIAL_CATEGORIES = ["入场培训", "复训", "签到单", "三级安全教育卡", "通知目录", "其他"]
ALERT_SEED_FILE = SEED_DIR / "01 防城港三期安全管理数据总台账.xlsx"
ALERT_DATA_FILE = DATA_DIR / "alert_台账.xlsx"

# Detail sheet column config: (sheet_name, category, type_name, date_col, dept_col, sub_col, problem_type_col)
# dept_col=0 / sub_col=0 / problem_type_col=0 means no such column in that sheet
DETAIL_SHEET_CONFIG = [
    ("工程公司挂牌督办单", "external", "挂牌督办", 2, 8, 11, 0),
    ("红黄牌", "external", "红黄牌", 3, 5, 12, 7),
    ("工程公司处理通报", "external", "处理通报", 4, 8, 10, 0),
    ("工程公司通报批评", "external", "处理通报", 5, 0, 11, 7),
    ("工程公司整改单", "external", "整改单", 6, 9, 12, 8),
    ("工程公司停工令", "external", "停工令", 5, 10, 13, 0),
    ("监理业主整改通知单", "external", "监理通知单", 5, 7, 10, 0),
    ("工程公司违章培训通知单", "external", "违章培训通知单", 7, 12, 3, 0),
    ("项目内部处理通报", "internal", "处理通报", 3, 7, 8, 6),
    ("项目整改通知单", "internal", "整改单", 3, 7, 8, 5),
    ("项目停工令", "internal", "停工令", 2, 6, 10, 0),
    ("项目违章培训通知单", "internal", "违章培训通知单", 7, 10, 4, 9),
]

NON_SUB_NAMES = {"分包", "责任分包", "项目总承包部", "总承包", "总承包部",
                  "综合车间", "钢结构队", "核岛一队", "水电队", "机械队", "搅拌站",
                  "驻场人员", "综合队", "测量队", "金属试验室", "机械设备管理部"}


def clean_sub_name(raw):
    """Normalize subcontractor name: remove | prefix, strip （...）suffix, skip invalid."""
    if not raw:
        return ""
    name = raw.strip()
    if name in NON_SUB_NAMES:
        return ""
    # Cells with many 、 are summary lists, not single subs (3+ names = 2+ separators)
    if name.count("、") >= 2:
        return ""
    # Take the part after the last |
    if "|" in name:
        name = name.rsplit("|", 1)[-1].strip()
    # Remove （...） / (...) suffix
    name = re.sub(r"[（(][^）)]*[）)]", "", name).strip()
    if not name or name in ("/", "None", "#N/A", "") or name in NON_SUB_NAMES:
        return ""
    return name


def clean_problem_type(raw):
    """Normalize problem type: remove leading numbers/tabs, split by comma, skip empty."""
    if not raw:
        return []
    text = str(raw).strip()
    # Remove leading digits, tabs, and separators like "1\t" or "12 "
    text = re.sub(r"^[\d\s\t]+", "", text)
    if not text or text in ("/", "None", "#N/A", "", "A", "B", "C", "A级", "B级", "C级"):
        return []
    # Split by comma or Chinese comma
    parts = re.split(r"[，,]", text)
    result = []
    for p in parts:
        p = p.strip()
        # Remove leading digits again for each part
        p = re.sub(r"^[\d\s\t]+", "", p).strip()
        if p and p not in ("/", "None", "#N/A", "", "A", "B", "C"):
            result.append(p)
    return result


# Alert dashboard data cache
_alert_data = None
DATA_DIR.mkdir(parents=True, exist_ok=True)
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)
TRAINING_LEDGER_DIR.mkdir(parents=True, exist_ok=True)
BRAKE_LEDGER_DIR.mkdir(parents=True, exist_ok=True)
FEEDBACK_DIR.mkdir(parents=True, exist_ok=True)

app = Flask(__name__, static_folder="public", static_url_path="")
app.config["MAX_CONTENT_LENGTH"] = 30 * 1024 * 1024

# Dashboard/statistics policy: D-level hazards remain searchable in the
# underlying ledger, but never contribute to dashboard totals or compliance.
NON_D_HAZARD_SQL = "UPPER(TRIM(COALESCE(hazard_level, ''))) NOT IN ('D', 'D级')"

REQUIRED_HEADERS = {"隐患单号", "检查人姓名", "检查日期", "检查单位"}
DATASET_TYPES = [
    {
        "id": "hazard_entry",
        "label": "隐患录入统计",
        "filenamePattern": r"^安全隐患信息表_[0-9]{6,20}(?:\s*\([0-9]+\))?\.xlsx$",
        "filenameExample": "安全隐患信息表_20260702141544.xlsx",
        "description": "更新隐患列表、人员录入统计及检查单位对比",
    },
]
DATE_HEADERS = {"检查日期", "整改期限", "实际整改日期", "实际验证日期", "关闭流程日期"}
DEFAULT_DEPARTMENTS = [
    "经理部", "安监部", "物资部", "技术部", "工程部", "质控部", "机械设备管理部",
    "核岛一队", "搅拌站", "水电队", "综合车间", "金属试验室", "钢结构队", "机械队", "测量队",
]
LEADER_DEPARTMENTS = {
    "经理部", "物资部", "技术部", "工程部", "质控部", "机械设备管理部",
    "核岛一队", "搅拌站", "水电队", "综合车间", "金属试验室", "钢结构队", "机械队", "测量队",
}
ROLE_RULES_VERSION = "2026-07-19-v4"
ROLE_STANDARDS = {
    "执行岗及以上": (5, "week"),
    "安全员": (10, "day"),
    "班组长": (2, "day"),
    "驻场代表": (3, "week"),
}
ROLE_ALIASES = {
    "执行岗及以上": "执行岗及以上",
    "执行岗以上": "执行岗及以上",
    "执行岗": "执行岗及以上",
    "安全员": "安全员",
    "安监部": "安全员",
    "班组长": "班组长",
    "班组": "班组长",
    "驻场代表": "驻场代表",
    "驻场": "驻场代表",
}
PEOPLE_HEADER_ALIASES = {
    "name": {"姓名", "人员", "人员姓名", "检查人姓名", "名字", "姓名/人员", "人员名称"},
    "department": {"部门", "单位", "所属部门", "部门/班组", "班组", "队伍", "组织", "科室"},
    "category": {"类型", "类别", "人员类型", "人员类别", "人员种类", "身份", "角色"},
}


def db():
    connection = sqlite3.connect(DB_PATH)
    connection.row_factory = sqlite3.Row
    connection.execute("PRAGMA foreign_keys=ON")
    return connection


def init_db():
    with db() as conn:
        conn.executescript("""
        CREATE TABLE IF NOT EXISTS hazards (
            hazard_no TEXT PRIMARY KEY,
            check_date TEXT NOT NULL,
            checker_raw TEXT,
            checker_name TEXT,
            check_unit TEXT,
            check_department TEXT,
            project_name TEXT,
            hazard_level TEXT,
            hazard_category TEXT,
            description TEXT,
            area TEXT,
            status TEXT,
            responsible_unit TEXT,
            responsible_department TEXT,
            responsible_team TEXT,
            raw_json TEXT NOT NULL,
            imported_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_hazards_date ON hazards(check_date);
        CREATE INDEX IF NOT EXISTS idx_hazards_checker ON hazards(checker_name);
        CREATE INDEX IF NOT EXISTS idx_hazards_unit ON hazards(check_unit);
        CREATE TABLE IF NOT EXISTS people (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            category TEXT NOT NULL,
            department TEXT NOT NULL DEFAULT '',
            target_count REAL NOT NULL DEFAULT 0,
            target_period TEXT NOT NULL DEFAULT 'week',
            active INTEGER NOT NULL DEFAULT 1
        );
        CREATE TABLE IF NOT EXISTS settings (key TEXT PRIMARY KEY, value TEXT NOT NULL);
        CREATE TABLE IF NOT EXISTS imports (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT NOT NULL,
            imported_at TEXT NOT NULL,
            row_count INTEGER NOT NULL,
            min_date TEXT,
            max_date TEXT
        );
        CREATE TABLE IF NOT EXISTS training_collections (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            description TEXT DEFAULT '',
            departments TEXT NOT NULL DEFAULT '',
            deadline TEXT DEFAULT '',
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS training_submissions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            collection_id INTEGER NOT NULL REFERENCES training_collections(id),
            department TEXT NOT NULL,
            sign_in_file TEXT DEFAULT '',
            photos TEXT DEFAULT '[]',
            notes TEXT DEFAULT '',
            submitted_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_submissions_collection ON training_submissions(collection_id);
        CREATE TABLE IF NOT EXISTS training_ledger_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            training_date TEXT NOT NULL,
            description TEXT NOT NULL DEFAULT '',
            schedule_time TEXT NOT NULL DEFAULT '19:30-21:00',
            schedule_period TEXT NOT NULL DEFAULT '晚上',
            training_location TEXT NOT NULL DEFAULT '',
            instructor TEXT NOT NULL DEFAULT '',
            audience TEXT NOT NULL DEFAULT '',
            participant_count INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS training_ledger_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id INTEGER NOT NULL REFERENCES training_ledger_events(id) ON DELETE CASCADE,
            original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL UNIQUE,
            display_name TEXT NOT NULL,
            kind TEXT NOT NULL DEFAULT 'file',
            content_type TEXT NOT NULL DEFAULT 'application/octet-stream',
            size INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_training_ledger_event ON training_ledger_files(event_id);
        CREATE TABLE IF NOT EXISTS subcontractor_attendance_dashboard (
            id INTEGER PRIMARY KEY CHECK (id=1),
            roster_json TEXT NOT NULL,
            result_json TEXT NOT NULL,
            roster_filename TEXT NOT NULL DEFAULT '',
            attendance_filename TEXT NOT NULL DEFAULT '',
            imported_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS training_categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            sort_order INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS brake_ledger_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            record_date TEXT NOT NULL,
            category TEXT NOT NULL DEFAULT '工程公司整改单',
            description TEXT NOT NULL DEFAULT '',
            issue_dept TEXT NOT NULL DEFAULT '',
            responsible_dept TEXT NOT NULL DEFAULT '',
            responsible_person TEXT NOT NULL DEFAULT '',
            area TEXT NOT NULL DEFAULT '',
            subcontractor TEXT NOT NULL DEFAULT '',
            team TEXT NOT NULL DEFAULT '',
            status TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS brake_ledger_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id INTEGER NOT NULL REFERENCES brake_ledger_events(id) ON DELETE CASCADE,
            original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL UNIQUE,
            display_name TEXT NOT NULL,
            kind TEXT NOT NULL DEFAULT 'file',
            content_type TEXT NOT NULL DEFAULT 'application/octet-stream',
            size INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_brake_ledger_event ON brake_ledger_files(event_id);
        CREATE TABLE IF NOT EXISTS brake_ledger_categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            sort_order INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS team_auth_dashboard (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            auth_filename TEXT NOT NULL DEFAULT '',
            result_json TEXT NOT NULL DEFAULT '{}',
            imported_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS feedback_events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            record_date TEXT NOT NULL,
            category TEXT NOT NULL DEFAULT '经验反馈',
            content TEXT NOT NULL DEFAULT '',
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS feedback_files (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            event_id INTEGER NOT NULL REFERENCES feedback_events(id) ON DELETE CASCADE,
            original_name TEXT NOT NULL,
            stored_name TEXT NOT NULL UNIQUE,
            display_name TEXT NOT NULL,
            kind TEXT NOT NULL DEFAULT 'file',
            content_type TEXT NOT NULL DEFAULT 'application/octet-stream',
            size INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_feedback_event ON feedback_files(event_id);
        CREATE TABLE IF NOT EXISTS feedback_categories (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            sort_order INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        """)
        ledger_columns = {row[1] for row in conn.execute("PRAGMA table_info(training_ledger_events)")}
        if "schedule_time" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN schedule_time TEXT NOT NULL DEFAULT '19:30-21:00'")
        if "schedule_period" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN schedule_period TEXT NOT NULL DEFAULT '晚上'")
        if "training_location" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN training_location TEXT NOT NULL DEFAULT ''")
        if "instructor" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN instructor TEXT NOT NULL DEFAULT ''")
        if "audience" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN audience TEXT NOT NULL DEFAULT ''")
        if "participant_count" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN participant_count INTEGER NOT NULL DEFAULT 0")
        if "category" not in ledger_columns:
            conn.execute("ALTER TABLE training_ledger_events ADD COLUMN category TEXT NOT NULL DEFAULT '专题培训'")
        conn.execute("""
            INSERT OR IGNORE INTO training_categories (name, sort_order, created_at)
            VALUES ('专题培训', 0, ?)
        """, (datetime.now().isoformat(timespec="seconds"),))
        conn.execute("""
            INSERT OR IGNORE INTO training_categories (name, sort_order, created_at)
            VALUES ('入场培训', 1, ?)
        """, (datetime.now().isoformat(timespec="seconds"),))
        brake_categories = [
            ("工程公司挂牌督办单", 0),
            ("红黄牌", 1),
            ("工程公司处理通报", 2),
            ("工程公司通报批评", 3),
            ("工程公司整改单", 4),
            ("工程公司停工令", 5),
            ("监理业主整改通知单", 6),
            ("行为偏差", 7),
            ("约谈记录", 8),
        ]
        for i, (name, order) in enumerate(brake_categories):
            conn.execute(
                "INSERT OR IGNORE INTO brake_ledger_categories (name, sort_order, created_at) VALUES (?, ?, ?)",
                (name, order, datetime.now().isoformat(timespec="seconds")),
            )
        brake_columns = {row[1] for row in conn.execute("PRAGMA table_info(brake_ledger_events)")}
        if "source_ref" not in brake_columns:
            conn.execute("ALTER TABLE brake_ledger_events ADD COLUMN source_ref TEXT NOT NULL DEFAULT ''")
            conn.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_brake_source_ref ON brake_ledger_events(source_ref) WHERE source_ref != ''")
        conn.execute(
            "INSERT OR IGNORE INTO feedback_categories (name, sort_order, created_at) VALUES ('经验反馈', 0, ?)",
            (datetime.now().isoformat(timespec="seconds"),),
        )
        fb_columns = {row[1] for row in conn.execute("PRAGMA table_info(feedback_events)")}
        if "participant_count" not in fb_columns:
            conn.execute("ALTER TABLE feedback_events ADD COLUMN participant_count INTEGER NOT NULL DEFAULT 0")
        if "source_ref" not in fb_columns:
            conn.execute("ALTER TABLE feedback_events ADD COLUMN source_ref TEXT NOT NULL DEFAULT ''")
            conn.execute("CREATE UNIQUE INDEX IF NOT EXISTS idx_feedback_source_ref ON feedback_events(source_ref) WHERE source_ref != ''")
        existing_ledger_files = conn.execute("""
            SELECT f.id,f.original_name,f.kind,e.name,e.training_date
            FROM training_ledger_files f
            JOIN training_ledger_events e ON e.id=f.event_id
        """).fetchall()
        for item in existing_ledger_files:
            suffix = Path(item["original_name"]).suffix.lower()
            year, month, day = map(int, item["training_date"].split("-"))
            file_label = "照片" if item["kind"] == "image" else "签到单"
            display_name = f"{year}年{month}月{day}日{item['name']}{file_label}{suffix}"
            conn.execute("UPDATE training_ledger_files SET display_name=? WHERE id=?", (display_name, item["id"]))
        defaults = {"internal_unit": "中建二局", "ratio_target": "5"}
        for key, value in defaults.items():
            conn.execute("INSERT OR IGNORE INTO settings(key,value) VALUES (?,?)", (key, value))


def clean_name(value):
    text = str(value or "").strip()
    return re.sub(r"^\[[^\]]+\]", "", text).strip()


def clean_header(value):
    return re.sub(r"[\s　：:（）()\[\]【】/\\]+", "", str(value or "")).strip()


def normalize_person_role(value):
    text = str(value or "").strip()
    compact = clean_header(text)
    for key, role in ROLE_ALIASES.items():
        if clean_header(key) in compact or compact in clean_header(key):
            return role
    return ""


def role_standard(role):
    return ROLE_STANDARDS.get(role, (0, "week"))


def find_people_sheet(workbook):
    alias_map = {field: {clean_header(x) for x in aliases} for field, aliases in PEOPLE_HEADER_ALIASES.items()}
    best = None
    for sheet in workbook.worksheets:
        max_rows = min(sheet.max_row, 30)
        max_cols = min(sheet.max_column, 30)
        for row_index in range(1, max_rows + 1):
            mapping = {}
            for col_index in range(1, max_cols + 1):
                header = clean_header(sheet.cell(row=row_index, column=col_index).value)
                if not header:
                    continue
                for field, aliases in alias_map.items():
                    if header in aliases and field not in mapping:
                        mapping[field] = col_index - 1
            score = len(mapping)
            if score >= 2 and ("name" in mapping and ("department" in mapping or "category" in mapping)):
                if score == 3:
                    return sheet, row_index, mapping
                if not best:
                    best = (sheet, row_index, mapping)
    if best:
        return best
    raise ValueError("未找到人员表头。至少需要识别“姓名”，并包含“部门”或“类型”；推荐表头：姓名、部门、类型。")


def parse_people_import(path):
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet, header_row, mapping = find_people_sheet(workbook)
    items = {}
    for row in sheet.iter_rows(min_row=header_row + 1, values_only=True):
        raw_name = row[mapping["name"]] if mapping.get("name", -1) < len(row) else ""
        name = clean_name(raw_name)
        if not name or name in {"姓名", "人员", "检查人姓名"}:
            continue
        department = ""
        if "department" in mapping and mapping["department"] < len(row):
            department = str(row[mapping["department"]] or "").strip()
        raw_role = row[mapping["category"]] if "category" in mapping and mapping["category"] < len(row) else ""
        if "班组长" in str(raw_role or "") and "驻场" in str(raw_role or ""):
            role = "驻场代表" if "驻场" in department else "班组长"
        else:
            role = normalize_person_role(raw_role)
        if not role:
            role = normalize_person_role(department)
        if not role:
            continue
        target, period = role_standard(role)
        items[name] = {
            "name": name,
            "category": role,
            "department": department,
            "target_count": target,
            "target_period": period,
            "active": 1,
        }
    if not items:
        raise ValueError("没有识别到可导入人员。请确认表中有“姓名、部门、类型”，类型为：安全员、执行岗及以上、班组长、驻场代表。")
    return list(items.values()), {"sheet": sheet.title, "headerRow": header_row, "columns": mapping}


def import_people(path, original_name, mode="append"):
    mode = "overwrite" if mode == "overwrite" else "append"
    items, source = parse_people_import(path)
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        if mode == "overwrite":
            conn.execute("DELETE FROM people")
        inserted = updated = 0
        for item in items:
            existing = conn.execute("SELECT id FROM people WHERE name=?", (item["name"],)).fetchone()
            values = (item["name"], item["category"], item["department"], item["target_count"], item["target_period"], item["active"])
            if existing:
                conn.execute("UPDATE people SET name=?,category=?,department=?,target_count=?,target_period=?,active=? WHERE id=?",
                             values + (existing["id"],))
                updated += 1
            else:
                conn.execute("INSERT INTO people(name,category,department,target_count,target_period,active) VALUES (?,?,?,?,?,?)", values)
                inserted += 1
        conn.execute("INSERT INTO imports(filename,imported_at,row_count,min_date,max_date) VALUES (?,?,?,?,?)",
                     (f"人员导入[{mode}]_{original_name}", now, len(items), None, None))
    return {
        "ok": True,
        "mode": mode,
        "count": len(items),
        "inserted": inserted,
        "updated": updated,
        "source": source,
        "filename": original_name,
    }


def detect_dataset_type(filename):
    name = Path(str(filename or "")).name
    for dataset in DATASET_TYPES:
        if re.fullmatch(dataset["filenamePattern"], name, flags=re.IGNORECASE):
            return dataset
    return None


def parse_training_date(value):
    text = str(value or "").strip()
    match = re.search(r"(\d{4})/(\d{1,2})/(\d{1,2})", text)
    if not match:
        return "", ""
    year, month, day = map(int, match.groups())
    date_text = date(year, month, day).isoformat()
    weekday_match = re.search(r"（([^）]+)）", text)
    return date_text, weekday_match.group(1) if weekday_match else ""


def split_training_items(value):
    text = str(value or "").strip()
    if not text:
        return []
    parts = []
    for line in re.split(r"[\n\r]+", text):
        for item in re.split(r"[、，,]", str(line).strip()):
            item = item.strip()
            if item:
                parts.append(item)
    return parts


def training_event(event_id, schedule_date, weekday, time_text, period, title, items=None, order=0):
    return {
        "id": event_id,
        "date": schedule_date,
        "weekday": weekday,
        "time": time_text,
        "period": period,
        "title": title,
        "items": items if items is not None else split_training_items(title),
        "order": order,
    }


def load_training_overrides():
    if not TRAINING_OVERRIDES_FILE.exists():
        return {}
    try:
        return json.loads(TRAINING_OVERRIDES_FILE.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}


def save_training_overrides(overrides):
    TRAINING_OVERRIDES_FILE.parent.mkdir(parents=True, exist_ok=True)
    TRAINING_OVERRIDES_FILE.write_text(json.dumps(overrides, ensure_ascii=False, indent=2), encoding="utf-8")


def load_training_materials():
    if not TRAINING_MATERIALS_FILE.exists():
        return []
    try:
        return json.loads(TRAINING_MATERIALS_FILE.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return []


def save_training_materials(materials):
    TRAINING_MATERIALS_FILE.parent.mkdir(parents=True, exist_ok=True)
    TRAINING_MATERIALS_FILE.write_text(json.dumps(materials, ensure_ascii=False, indent=2), encoding="utf-8")


def apply_training_overrides(events):
    overrides = load_training_overrides()
    for event in events:
        override = overrides.get(event["id"])
        if not override:
            continue
        for key in ("time", "period", "title", "items"):
            if key in override:
                event[key] = override[key]
        event["edited"] = True
    return events


def add_annual_retraining(events):
    plans = {
        "2026-07-04": "7月年度复训",
        "2026-07-05": "7月年度复训",
        "2026-07-24": "8月年度复训",
        "2026-07-25": "8月年度复训",
    }
    weekday_names = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    for schedule_date, label in plans.items():
        current = date.fromisoformat(schedule_date)
        weekday = weekday_names[current.weekday()]
        events.append(training_event(
            f"annual:{schedule_date}:signin", schedule_date, weekday, "08:00-10:00",
            "年度复训签到", f"{label}公司级签到", ["年度复训", "公司级签到"], 80,
        ))
        events.append(training_event(
            f"annual:{schedule_date}:exam", schedule_date, weekday, "10:00-11:30",
            "年度复训考试", f"11:00 {label}公司级考试", ["年度复训", "公司级考试"], 81,
        ))
        events.append(training_event(
            f"annual:{schedule_date}:night", schedule_date, weekday, "19:00-20:30",
            "晚上", label, [label], 90,
        ))
    return events


# ── August Training Schedule ──────────────────────────────────────────

AUGUST_SCHEDULE_FILE = DATA_DIR / "august_schedule.json"


def _generate_august_schedule():
    """Generate August 2026 schedule: Mon/Wed/Fri=公司级, Tue/Thu/Sat=项目级, Sun off."""
    import calendar as _cal
    days = []
    for d in range(1, 32):
        wd = _cal.weekday(2026, 8, d)  # 0=Mon..6=Sun
        if wd == 6:  # Sunday
            continue
        level = "公司级" if wd in (0, 2, 4) else "项目级"
        weekdays_cn = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
        days.append({
            "date": f"2026-08-{d:02d}",
            "weekday": weekdays_cn[wd],
            "level": level,
            "title": "",
            "time": "19:00-20:30",
            "location": "",
            "instructor": "",
            "note": "",
        })
    return {"year": 2026, "month": 8, "days": days, "week_overrides": {}}


def _load_august_schedule():
    if AUGUST_SCHEDULE_FILE.exists():
        try:
            return json.loads(AUGUST_SCHEDULE_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    data = _generate_august_schedule()
    _save_august_schedule(data)
    return data


def _save_august_schedule(data):
    AUGUST_SCHEDULE_FILE.parent.mkdir(parents=True, exist_ok=True)
    AUGUST_SCHEDULE_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


@app.get("/api/training/august-schedule")
def august_schedule_get():
    return jsonify(_load_august_schedule())


@app.post("/api/training/august-schedule/update-day")
def august_schedule_update_day():
    if not training_password_ok():
        return jsonify(error="编辑密码错误"), 403
    body = request.get_json(force=True) or {}
    date_str = str(body.get("date") or "").strip()
    if not date_str:
        return jsonify(error="缺少日期"), 400
    data = _load_august_schedule()
    for d in data["days"]:
        if d["date"] == date_str:
            for k in ("title", "time", "location", "instructor", "note", "level"):
                if k in body:
                    d[k] = str(body[k] or "").strip()
            _save_august_schedule(data)
            return jsonify(ok=True, day=d)
    return jsonify(error="日期未找到"), 404


@app.post("/api/training/august-schedule/override-week")
def august_schedule_override_week():
    """Override an entire week: {week_start: '2026-08-03', days: [{weekday,level,title,...}]}"""
    if not training_password_ok():
        return jsonify(error="编辑密码错误"), 403
    body = request.get_json(force=True) or {}
    week_start = str(body.get("week_start") or "").strip()
    custom_days = body.get("days") or []
    if not week_start or not custom_days:
        return jsonify(error="缺少 week_start 或 days"), 400
    data = _load_august_schedule()
    overrides = data.setdefault("week_overrides", {})
    overrides[week_start] = custom_days
    _save_august_schedule(data)
    return jsonify(ok=True)


@app.post("/api/training/august-schedule/reset-week")
def august_schedule_reset_week():
    """Reset a week back to the default pattern."""
    if not training_password_ok():
        return jsonify(error="编辑密码错误"), 403
    body = request.get_json(force=True) or {}
    week_start = str(body.get("week_start") or "").strip()
    if not week_start:
        return jsonify(error="缺少 week_start"), 400
    data = _load_august_schedule()
    data.get("week_overrides", {}).pop(week_start, None)
    _save_august_schedule(data)
    return jsonify(ok=True)


def training_password_ok():
    body = request.get_json(force=True) or {}
    pwd = str(body.get("password") or request.args.get("password", ""))
    return pwd == TRAINING_EDIT_PASSWORD


# ── Training Schedule (legacy Excel-based) ────────────────────────────

def read_training_schedule():
    events = []
    if TRAINING_SCHEDULE_FILE.exists():
        workbook = load_workbook(TRAINING_SCHEDULE_FILE, read_only=True, data_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        for row in sheet.iter_rows(min_row=3, values_only=True):
            schedule_date, weekday = parse_training_date(row[0] if len(row) > 0 else "")
            if not schedule_date:
                continue
            day_index = len([event for event in events if event["date"] == schedule_date])
            morning_text = row[1] if len(row) > 1 else None
            morning_lines = [line.strip() for line in re.split(r"[\n\r]+", str(morning_text or "").strip()) if line.strip()]
            if morning_lines:
                events.append(training_event(
                    f"base:{schedule_date}:morning1", schedule_date, weekday, "08:00-10:00",
                    "上午第一场", morning_lines[0], order=10 + day_index,
                ))
            if len(morning_lines) > 1:
                second_text = "\n".join(morning_lines[1:])
                events.append(training_event(
                    f"base:{schedule_date}:morning2", schedule_date, weekday, "10:00-11:30",
                    "上午第二场", second_text, order=20 + day_index,
                ))
            afternoon_parts = []
            for col_index in (3, 4):
                if len(row) > col_index and row[col_index]:
                    afternoon_parts.append(str(row[col_index]).strip())
            if afternoon_parts:
                text = "、".join(afternoon_parts)
                events.append(training_event(
                    f"base:{schedule_date}:afternoon", schedule_date, weekday, "14:00-17:30",
                    "下午", text, order=30 + day_index,
                ))
            if len(row) > 5 and row[5]:
                text = str(row[5]).strip()
                events.append(training_event(
                    f"base:{schedule_date}:night", schedule_date, weekday, "19:00-20:30",
                    "晚上", text, order=90 + day_index,
                ))
        workbook.close()
    add_annual_retraining(events)
    apply_training_overrides(events)
    weekday_names = ["周一", "周二", "周三", "周四", "周五", "周六", "周日"]
    with db() as conn:
        linked_events = conn.execute("""
            SELECT id,name,training_date,description,schedule_time,schedule_period,
                   training_location,instructor,audience,participant_count,category
            FROM training_ledger_events
        """).fetchall()
        linked_ids = [row["id"] for row in linked_events]
        files_by_event = {}
        if linked_ids:
            placeholders = ",".join("?" for _ in linked_ids)
            file_rows = conn.execute(f"""
                SELECT id,event_id,display_name,kind,content_type,size
                FROM training_ledger_files
                WHERE event_id IN ({placeholders})
                ORDER BY id
            """, linked_ids).fetchall()
            for frow in file_rows:
                fdict = dict(frow)
                fdict["download_url"] = f"/api/training-ledger/files/{fdict['id']}/download"
                fdict["preview_url"] = f"/api/training-ledger/files/{fdict['id']}/preview"
                files_by_event.setdefault(frow["event_id"], []).append(fdict)
    for row in linked_events:
        current = date.fromisoformat(row["training_date"])
        event = training_event(
            f"ledger:{row['id']}", row["training_date"], weekday_names[current.weekday()],
            row["schedule_time"], row["schedule_period"], row["name"], [row["name"]], 95,
        )
        event["linked"] = True
        event["description"] = row["description"]
        event["training_location"] = row["training_location"]
        event["instructor"] = row["instructor"]
        event["audience"] = row["audience"]
        event["participant_count"] = row["participant_count"]
        event["category"] = row["category"] or ""
        event["files"] = files_by_event.get(row["id"], [])
        events.append(event)
    return sorted(events, key=lambda item: (item["date"], item.get("order", 0), item["time"], item["period"]))


def week_range_for(value):
    current = iso_date(value) if value else date.today().isoformat()
    current_date = datetime.strptime(current, "%Y-%m-%d").date()
    start = current_date - timedelta(days=current_date.weekday())
    end = start + timedelta(days=6)
    return start.isoformat(), end.isoformat()


def iso_date(value):
    if value is None or value == "":
        return ""
    if isinstance(value, datetime):
        return value.date().isoformat()
    if isinstance(value, date):
        return value.isoformat()
    if isinstance(value, (int, float)):
        return (datetime(1899, 12, 30) + timedelta(days=float(value))).date().isoformat()
    text = str(value).strip()
    for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"):
        try:
            return datetime.strptime(text[:19], fmt).date().isoformat()
        except ValueError:
            pass
    return text[:10]


def json_value(value):
    if isinstance(value, (datetime, date)):
        return value.isoformat(sep=" ") if isinstance(value, datetime) else value.isoformat()
    return value


def find_data_sheet(workbook):
    for sheet in workbook.worksheets:
        header = [cell.value for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
        if REQUIRED_HEADERS.issubset(set(header)):
            return sheet, header
    raise ValueError("未找到完整表头，需要包含：隐患单号、检查人姓名、检查日期、检查单位")


def import_hazards(path, original_name):
    workbook = load_workbook(path, read_only=True, data_only=True)
    sheet, headers = find_data_sheet(workbook)
    now = datetime.now().isoformat(timespec="seconds")
    records = []
    for row in sheet.iter_rows(min_row=2, values_only=True):
        item = {str(headers[i]).strip(): json_value(value) for i, value in enumerate(row) if i < len(headers) and headers[i]}
        hazard_no = str(item.get("隐患单号") or "").strip()
        if not hazard_no:
            continue
        check_date = iso_date(item.get("检查日期"))
        records.append((
            hazard_no, check_date, str(item.get("检查人姓名") or ""), clean_name(item.get("检查人姓名")),
            str(item.get("检查单位") or "").strip(), str(item.get("检查部门") or "").strip(),
            str(item.get("项目名称") or "").strip(), str(item.get("隐患级别") or "").strip(),
            str(item.get("隐患分类") or "").strip(), str(item.get("隐患描述") or "").strip(),
            str(item.get("区域") or "").strip(), str(item.get("流程状态") or item.get("状态") or "").strip(),
            str(item.get("责任单位") or "").strip(), str(item.get("责任部门") or "").strip(),
            str(item.get("责任班组") or "").strip(), json.dumps(item, ensure_ascii=False), now,
        ))
    if not records:
        raise ValueError("表格中没有可导入的隐患记录")
    with db() as conn:
        conn.executemany("""INSERT OR REPLACE INTO hazards VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)""", records)
        dates = [r[1] for r in records if r[1]]
        conn.execute("INSERT INTO imports(filename,imported_at,row_count,min_date,max_date) VALUES (?,?,?,?,?)",
                     (original_name, now, len(records), min(dates) if dates else None, max(dates) if dates else None))
    return {"count": len(records), "minDate": min(dates), "maxDate": max(dates), "filename": original_name,
            "datasetType": "hazard_entry", "datasetLabel": "隐患录入统计"}


def seed_people():
    legacy = SEED_DIR / "隐患统计查询工具1.02.xlsx"
    if not legacy.exists():
        return
    with db() as conn:
        if conn.execute("SELECT COUNT(*) FROM people").fetchone()[0]:
            return
    workbook = load_workbook(legacy, read_only=True, data_only=True)
    sheet = workbook["数据查询界面"]
    found = {}
    for row in sheet.iter_rows(min_row=8, max_row=196, values_only=True):
        for category, name, target in ((row[1], row[2], row[3]), (row[8], row[9], row[10])):
            if not name or not category or str(category).startswith("统计"):
                continue
            name = clean_name(name)
            category = str(category).strip()
            if not name or name in found:
                continue
            period = "day" if category == "安监部" else "week"
            default_target = 10 if category == "安监部" else (7 if "班组长" in category else 5)
            numeric_target = float(target) if isinstance(target, (int, float)) else default_target
            found[name] = (name, category, category if category in DEFAULT_DEPARTMENTS else "", numeric_target, period)
    with db() as conn:
        conn.executemany("INSERT OR IGNORE INTO people(name,category,department,target_count,target_period) VALUES (?,?,?,?,?)", found.values())


def apply_role_rules():
    with db() as conn:
        current = conn.execute("SELECT value FROM settings WHERE key='role_rules_version'").fetchone()
        if current and current[0] == ROLE_RULES_VERSION:
            return
        people = conn.execute("SELECT id,name,category,department FROM people").fetchall()
        for person in people:
            legacy_category = str(person["category"] or "").strip()
            department = str(person["department"] or "").strip()
            role = None
            if person["name"] == "赵强强":
                role = "执行岗及以上"
            elif legacy_category in ("安全员", "安监部") or department == "安监部":
                role = "安全员"
                department = department or "安监部"
            elif "班组长" in legacy_category:
                role = "班组长"
            elif legacy_category == "驻场代表":
                role = "驻场代表"
            elif legacy_category in ("执行岗及以上", "执行岗以上"):
                role = None
            else:
                role = normalize_person_role(legacy_category)
            if role:
                target, period = ROLE_STANDARDS[role]
                conn.execute("UPDATE people SET category=?,department=?,target_count=?,target_period=?,active=1 WHERE id=?",
                             (role, department, target, period, person["id"]))
            else:
                conn.execute("UPDATE people SET active=0 WHERE id=?", (person["id"],))
        conn.execute("INSERT INTO settings(key,value) VALUES('role_rules_version',?) ON CONFLICT(key) DO UPDATE SET value=excluded.value",
                     (ROLE_RULES_VERSION,))


def load_alert_data():
    """Parse the 总台账 Excel. Tries DATA_DIR first, falls back to SEED_DIR."""
    global _alert_data
    if _alert_data is not None:
        return _alert_data

    src = None
    if ALERT_DATA_FILE.exists():
        src = ALERT_DATA_FILE
    elif ALERT_SEED_FILE.exists():
        src = ALERT_SEED_FILE

    if not src:
        _alert_data = {}
        return _alert_data

    wb = load_workbook(src, data_only=True)

    def cell(ws, r, c):
        v = ws.cell(row=r, column=c).value
        return v if v is not None else 0

    # --- Parse 总台账及统计分析 ---
    ws = wb["总台账及统计分析"]
    MONTHS = ["1月", "2月", "3月", "4月", "5月", "6月", "7月", "8月", "9月", "10月", "11月", "12月"]

    external_types = [
        ("挂牌督办", 4), ("管理约谈", 5), ("红黄牌", 6), ("处理通报", 7),
        ("停工令", 9), ("整改单", 11), ("违章培训通知单", 13), ("监理通知单", 16),
    ]
    internal_types = [
        ("停工令", 10), ("处理通报", 8), ("整改单", 12), ("违章培训通知单", 14),
    ]

    def parse_monthly(row):
        monthly = []
        for c in range(2, 14):  # cols B-M
            v = cell(ws, row, c)
            monthly.append(int(v) if v else 0)
        return monthly

    external = []
    for name, row in external_types:
        monthly = parse_monthly(row)
        external.append({"name": name, "monthly": monthly, "total": sum(monthly)})

    internal = []
    for name, row in internal_types:
        monthly = parse_monthly(row)
        internal.append({"name": name, "monthly": monthly, "total": sum(monthly)})

    # Department data (rows 5-14, columns Q-AG)
    # Dept col mapping: R/S=整改单(当月/累计), T/U=监理, V/W=项目内部,
    #   X/Y=违章培训(当月/累计), Z/AA=项目内部违章,
    #   AB/AC=处理通报(当月/累计), AD/AE=项目内部通报,
    #   AF=黄牌, AG=挂牌督办
    dept_names = []
    dept_external_rect = []  # 整改单 累计
    dept_external_violation = []  # 违章培训 累计
    dept_external_notice = []  # 处理通报 累计
    dept_yellow = []  # 黄牌
    dept_supervision = []  # 挂牌督办

    for r in range(5, 15):
        name = str(cell(ws, r, 17) or "")  # Q column
        if not name or name == "0":
            continue
        dept_names.append(name)
        dept_external_rect.append(int(cell(ws, r, 19)))  # S=累计
        dept_external_violation.append(int(cell(ws, r, 25)))  # Y=累计
        dept_external_notice.append(int(cell(ws, r, 29)))  # AC=累计
        dept_yellow.append(int(cell(ws, r, 32)))  # AF
        dept_supervision.append(int(cell(ws, r, 33)))  # AG

    departments = {
        "names": dept_names,
        "external": [
            {"label": "整改单", "values": dept_external_rect},
            {"label": "违章培训通知单", "values": dept_external_violation},
            {"label": "处理通报", "values": dept_external_notice},
            {"label": "黄牌", "values": dept_yellow},
            {"label": "挂牌督办", "values": dept_supervision},
        ],
        # Internal dept data: project rectification + project violation + project notice
        # From cols: V=当月, W=累计(整改单), Z=当月, AA=累计(违章), AD=当月, AE=累计(通报)
        "internal": [],
    }
    for r in range(5, 15):
        name = str(cell(ws, r, 17) or "")
        if not name or name == "0":
            continue
        departments["internal"].append({
            "name": name,
            "rectification": int(cell(ws, r, 23)),  # W=累计
            "violation": int(cell(ws, r, 27)),  # AA=累计
            "notice": int(cell(ws, r, 31)),  # AE=累计
        })

    # --- Parse subcontractor data from detail sheets ---
    subcontractors = {}
    detail_sheets = [
        "工程公司挂牌督办单", "红黄牌", "工程公司处理通报", "工程公司通报批评",
        "工程公司整改单", "工程公司停工令", "监理业主整改通知单",
        "项目内部处理通报", "项目整改通知单", "项目停工令",
        "项目违章培训通知单", "工程公司违章培训通知单",
    ]
    for sheet_name in detail_sheets:
        if sheet_name not in wb.sheetnames:
            continue
        dws = wb[sheet_name]
        # Determine column index for subcontractor name (varies by sheet)
        header_row = 1
        sub_col = None
        for c in range(1, dws.max_column + 1):
            h = str(dws.cell(row=1, column=c).value or "")
            if "分包" in h:
                sub_col = c
                break
        if sub_col is None:
            continue
        for r in range(2, dws.max_row + 1):
            name = clean_sub_name(str(dws.cell(row=r, column=sub_col).value or ""))
            if name:
                subcontractors[name] = subcontractors.get(name, 0) + 1

    sub_list = sorted(
        [{"name": k, "count": v} for k, v in subcontractors.items()],
        key=lambda x: -x["count"],
    )

    # --- Parse detail records for monthly filtering & chart data ---
    sub_monthly = {}  # {month: {sub_name: {"external": {type: cnt}, "internal": {type: cnt}}}}
    dept_monthly = {}  # {month: {dept_name: {"external": {type: cnt}, "internal": {type: cnt}}}}
    sub_totals = {}  # {sub_name: {"external": total, "internal": total}}
    dept_totals = {}  # {dept_name: {"external": total, "internal": total}}
    detail_records = []  # [{date, sub_name, dept_name, category, type_name}]

    for sheet_name, category, type_name, date_col, dept_col, sub_col, prob_col in DETAIL_SHEET_CONFIG:
        if sheet_name not in wb.sheetnames:
            continue
        dws = wb[sheet_name]
        for r in range(2, dws.max_row + 1):
            date_val = dws.cell(row=r, column=date_col).value
            if date_val is None:
                continue
            parsed = iso_date(date_val)
            if not parsed:
                continue
            try:
                month = int(parsed[5:7])
                if month < 1 or month > 12:
                    continue
            except (ValueError, IndexError):
                continue

            sub_name = ""
            if sub_col > 0:
                sub_name = clean_sub_name(str(dws.cell(row=r, column=sub_col).value or ""))

            dept_name = ""
            if dept_col > 0:
                dept_name = str(dws.cell(row=r, column=dept_col).value or "").strip()
                if dept_name in ("/", "None", "#N/A", ""):
                    dept_name = ""

            # Parse problem type(s) from the detail sheet
            problem_types = []
            if prob_col > 0:
                pv = dws.cell(row=r, column=prob_col).value
                if pv:
                    problem_types = clean_problem_type(str(pv))

            for pt in problem_types:
                detail_records.append({
                    "date": parsed, "sub_name": sub_name,
                    "dept_name": dept_name, "category": category,
                    "type_name": type_name, "problem_type": pt,
                })
            if not problem_types:
                detail_records.append({
                    "date": parsed, "sub_name": sub_name,
                    "dept_name": dept_name, "category": category,
                    "type_name": type_name, "problem_type": "",
                })

            if sub_name:
                m = sub_monthly.setdefault(month, {})
                s = m.setdefault(sub_name, {"external": {}, "internal": {}})
                s[category][type_name] = s[category].get(type_name, 0) + 1
                st = sub_totals.setdefault(sub_name, {"external": 0, "internal": 0})
                st[category] += 1

            if dept_name:
                m = dept_monthly.setdefault(month, {})
                d = m.setdefault(dept_name, {"external": {}, "internal": {}})
                d[category][type_name] = d[category].get(type_name, 0) + 1
                dt = dept_totals.setdefault(dept_name, {"external": 0, "internal": 0})
                dt[category] += 1

    # --- Parse scores ---
    scores = {"star5_2025": [], "aqhb_2025": [], "star5_2026": [], "aqhb_2026": []}
    if "五星评估、安质环考核得分" in wb.sheetnames:
        sws = wb["五星评估、安质环考核得分"]
        # 2025 data rows 14-16
        for i, m in enumerate(MONTHS):
            v = sws.cell(row=15, column=2 + i).value
            if v is not None:
                scores["star5_2025"].append({"month": m, "score": float(v)})
        for i, m in enumerate(MONTHS):
            v = sws.cell(row=16, column=2 + i).value
            if v is not None:
                scores["aqhb_2025"].append({"month": m, "score": float(v)})
        # 2026 data rows 78-79
        for i, m in enumerate(MONTHS):
            v = sws.cell(row=78, column=2 + i).value
            if v is not None:
                scores["star5_2026"].append({"month": m, "score": float(v)})
        for i, m in enumerate(MONTHS):
            v = sws.cell(row=79, column=2 + i).value
            if v is not None:
                scores["aqhb_2026"].append({"month": m, "score": float(v)})

    wb.close()

    _alert_data = {
        "external": external,
        "internal": internal,
        "departments": departments,
        "subcontractors": sub_list,
        "scores": scores,
        "months": MONTHS,
        "sub_monthly": sub_monthly,
        "dept_monthly": dept_monthly,
        "sub_totals": sub_totals,
        "dept_totals": dept_totals,
        "detail_records": detail_records,
    }
    return _alert_data


def seed_data():
    with db() as conn:
        count = conn.execute("SELECT COUNT(*) FROM hazards").fetchone()[0]
    if count:
        return
    candidates = sorted(SEED_DIR.glob("安全隐患信息表*.xlsx"))
    if candidates:
        import_hazards(candidates[-1], candidates[-1].name)


def range_args():
    end = request.args.get("end") or date.today().isoformat()
    start = request.args.get("start") or (date.fromisoformat(end) - timedelta(days=6)).isoformat()
    return start, end


@app.get("/")
def index():
    return send_from_directory(app.static_folder, "index.html")


@app.get("/training")
def training_page():
    return send_from_directory(app.static_folder, "training.html")


@app.get("/train")
def train_page():
    return send_from_directory(app.static_folder, "training.html")


@app.get("/training-ledger")
def training_ledger_page():
    return send_from_directory(app.static_folder, "training-ledger.html")


@app.get("/subcontractor-attendance")
def subcontractor_attendance_page():
    return send_from_directory(app.static_folder, "subcontractor-attendance.html")


@app.get("/api/health")
def health():
    with db() as conn:
        count = conn.execute("SELECT COUNT(*) FROM hazards").fetchone()[0]
    return jsonify(ok=True, records=count)


def workbook_rows(upload_file):
    workbook = load_workbook(upload_file.stream, read_only=True, data_only=True)
    try:
        for sheet in workbook.worksheets:
            sheet.reset_dimensions()
            yield sheet.title, sheet.iter_rows(values_only=True)
    finally:
        workbook.close()


def parse_subcontractor_roster(upload_file):
    people = []
    seen = set()
    for _, rows in workbook_rows(upload_file):
        for row in rows:
            values = [str(value or "").strip() for value in row]
            nonempty = [value for value in values if value]
            if not nonempty:
                continue
            name = nonempty[0]
            if clean_header(name) in {"姓名", "分包代表", "代表姓名", "序号"}:
                continue
            name = re.sub(r"^\d+[.、\s]+", "", name).strip()
            if name and name not in seen:
                seen.add(name)
                role = nonempty[1] if len(nonempty) > 1 else ""
                people.append({"name": name, "role": role})
        if people:
            break
    return people


def parse_attendance_rows(upload_file):
    records = []
    for sheet_name, rows in workbook_rows(upload_file):
        buffered = list(rows)
        header_index = name_index = date_index = None
        for index, row in enumerate(buffered[:12]):
            headers = [clean_header(value) for value in row]
            for candidate in ("姓名", "成员姓名", "人员姓名"):
                if candidate in headers:
                    name_index = headers.index(candidate)
                    break
            for candidate in ("日期", "签到日期", "打卡日期"):
                if candidate in headers:
                    date_index = headers.index(candidate)
                    break
            if name_index is not None and date_index is not None:
                header_index = index
                break
        if header_index is None:
            continue
        for row in buffered[header_index + 1:]:
            if len(row) <= max(name_index, date_index):
                continue
            name = str(row[name_index] or "").strip()
            date_text = iso_date(row[date_index])
            if name and re.fullmatch(r"\d{4}-\d{2}-\d{2}", date_text):
                records.append({"name": name, "date": date_text, "sheet": sheet_name})
        if records:
            break
    return records


@app.post("/api/subcontractor-attendance/analyze")
def subcontractor_attendance_analyze():
    roster_file = request.files.get("roster")
    attendance_file = request.files.get("attendance")
    if not attendance_file or not attendance_file.filename:
        return jsonify(error="请选择签到台账"), 400
    if roster_file and roster_file.filename and not roster_file.filename.lower().endswith(".xlsx"):
        return jsonify(error="名单文件目前只支持 .xlsx"), 400
    if not attendance_file.filename.lower().endswith(".xlsx"):
        return jsonify(error="目前只支持 .xlsx 文件"), 400
    try:
        if roster_file and roster_file.filename:
            roster = parse_subcontractor_roster(roster_file)
            roster_filename = roster_file.filename
        else:
            with db() as conn:
                saved = conn.execute(
                    "SELECT roster_json,roster_filename FROM subcontractor_attendance_dashboard WHERE id=1"
                ).fetchone()
            if not saved:
                return jsonify(error="首次统计请先导入分包代表名单"), 400
            roster = json.loads(saved["roster_json"])
            roster_filename = saved["roster_filename"]
        attendance = parse_attendance_rows(attendance_file)
    except Exception as exc:
        return jsonify(error=f"Excel 读取失败：{exc}"), 400
    if not roster:
        return jsonify(error="名单表中未识别到姓名"), 400
    if isinstance(roster[0], str):
        roster = [{"name": n, "role": ""} for n in roster]
    roster_names = [item["name"] for item in roster]
    if not attendance:
        return jsonify(error="签到台账中未识别到“姓名”和“日期”明细"), 400
    all_dates = sorted({item["date"] for item in attendance})
    start, end = all_dates[0], all_dates[-1]
    period_days = (date.fromisoformat(end) - date.fromisoformat(start)).days + 1
    daily_names = {day: set() for day in all_dates}
    person_dates = {}
    person_punches = {}
    for item in attendance:
        daily_names[item["date"]].add(item["name"])
        person_dates.setdefault(item["name"], set()).add(item["date"])
        person_punches[item["name"]] = person_punches.get(item["name"], 0) + 1
    people = []
    for entry in roster:
        name = entry["name"]
        dates = sorted(person_dates.get(name, set()))
        signed_days = len(dates)
        people.append({
            "name": name,
            "role": entry["role"],
            "signed_days": signed_days,
            "missing_days": max(0, period_days - signed_days),
            "attendance_rate": round(signed_days / period_days, 4) if period_days else 0,
            "punch_count": person_punches.get(name, 0),
            "first_date": dates[0] if dates else "",
            "last_date": dates[-1] if dates else "",
            "dates": dates,
            "status": "全勤" if signed_days == period_days else ("未签到" if signed_days == 0 else "部分签到"),
        })
    people.sort(key=lambda item: (item["signed_days"], item["name"]))
    roster_set = set(roster_names)
    unmatched = sorted({item["name"] for item in attendance if item["name"] not in roster_set})
    daily = [{
        "date": day,
        "signed_count": len(daily_names[day] & roster_set),
        "missing_count": len(roster) - len(daily_names[day] & roster_set),
        "rate": round(len(daily_names[day] & roster_set) / len(roster), 4),
    } for day in all_dates]
    total_signed_days = sum(item["signed_days"] for item in people)
    imported_at = datetime.now().isoformat(timespec="seconds")
    result = {
        "start": start,
        "end": end,
        "period_days": period_days,
        "roster_count": len(roster),
        "attendance_record_count": len(attendance),
        "full_attendance_count": sum(item["status"] == "全勤" for item in people),
        "zero_attendance_count": sum(item["status"] == "未签到" for item in people),
        "average_days": round(total_signed_days / len(roster), 2),
        "overall_rate": round(total_signed_days / (len(roster) * period_days), 4) if period_days else 0,
        "people": people,
        "daily": daily,
        "unmatched_names": unmatched,
        "roster_filename": roster_filename,
        "attendance_filename": attendance_file.filename,
        "imported_at": imported_at,
    }
    with db() as conn:
        conn.execute("""
            INSERT INTO subcontractor_attendance_dashboard
            (id,roster_json,result_json,roster_filename,attendance_filename,imported_at)
            VALUES (1,?,?,?,?,?)
            ON CONFLICT(id) DO UPDATE SET
                roster_json=excluded.roster_json,
                result_json=excluded.result_json,
                roster_filename=excluded.roster_filename,
                attendance_filename=excluded.attendance_filename,
                imported_at=excluded.imported_at
        """, (
            json.dumps(roster, ensure_ascii=False),
            json.dumps(result, ensure_ascii=False),
            roster_filename,
            attendance_file.filename,
            imported_at,
        ))
    return jsonify(result)


@app.get("/api/subcontractor-attendance/current")
def subcontractor_attendance_current():
    with db() as conn:
        saved = conn.execute(
            "SELECT result_json FROM subcontractor_attendance_dashboard WHERE id=1"
        ).fetchone()
    if not saved:
        return jsonify(error="尚未导入签到统计数据"), 404
    return jsonify(json.loads(saved["result_json"]))


@app.get("/api/subcontractor-attendance/export-image")
def subcontractor_attendance_export_image():
    with db() as conn:
        saved = conn.execute(
            "SELECT result_json FROM subcontractor_attendance_dashboard WHERE id=1"
        ).fetchone()
    if not saved:
        return jsonify(error="尚未导入签到统计数据"), 404
    data = json.loads(saved["result_json"])
    role_filter = (request.args.get("role") or "").strip()
    people = data["people"]
    if role_filter:
        people = [p for p in people if p.get("role") == role_filter]
    return _generate_attendance_image(data, people, role_filter or "全部")


def _generate_attendance_image(data, people, title_role):
    from PIL import Image, ImageDraw, ImageFont
    header_h = 150
    row_h = 36
    margin = 48
    col_widths = [110, 90, 84, 92, 80, 140]
    col_labels = ["姓名", "角色", "签到天数", "未签到天数", "签到率", "状态"]
    table_w = sum(col_widths)
    img_w = table_w + margin * 2
    img_h = header_h + row_h + len(people) * row_h + 60
    img = Image.new("RGB", (img_w, img_h), "#e6f7f0")
    draw = ImageDraw.Draw(img)
    try:
        font_title = ImageFont.truetype("msyh.ttc", 40)
        font_sub = ImageFont.truetype("msyh.ttc", 18)
        font_header = ImageFont.truetype("msyhbd.ttf", 16)
        font_row = ImageFont.truetype("msyh.ttc", 15)
        font_row_bold = ImageFont.truetype("msyhbd.ttf", 15)
    except (OSError, IOError):
        font_title = ImageFont.load_default()
        font_sub = font_title
        font_header = font_title
        font_row = font_title
        font_row_bold = font_title
    # header gradient
    for y in range(header_h):
        ratio = y / header_h
        r = int(0x0A + (0x0D - 0x0A) * ratio)
        g = int(0x2E + (0x7A - 0x2E) * ratio)
        b = int(0x26 + (0x5C - 0x26) * ratio)
        draw.line([(0, y), (img_w, y)], fill=(r, g, b))
    draw.text((margin, 28), "SUBCONTRACTOR ATTENDANCE", fill=(0x75, 0xD9, 0xBB), font=font_sub)
    draw.text((margin, 68), "分包代表签到看板", fill=(255, 255, 255), font=font_title)
    subtitle = f"{data['start']} 至 {data['end']} | {title_role} | {len(people)} 人"
    draw.text((margin, 118), subtitle, fill=(0xD8, 0xEB, 0xE5), font=font_sub)
    # table header
    x = margin
    header_y = header_h
    draw.rectangle([margin, header_y, img_w - margin, header_y + row_h], fill="#0d7a5c")
    for i, (label, w) in enumerate(zip(col_labels, col_widths)):
        draw.text((x + 12, header_y + 8), label, fill=(255, 255, 255), font=font_header)
        x += w
    # rows
    for row_i, person in enumerate(people):
        y = header_y + row_h + row_i * row_h
        bg = (255, 255, 255) if row_i % 2 == 0 else (0xF5, 0xF9, 0xF7)
        draw.rectangle([margin, y, img_w - margin, y + row_h], fill=bg)
        draw.line([(margin, y + row_h), (img_w - margin, y + row_h)], fill=(0xDC, 0xE8, 0xE3))
        values = [
            person["name"],
            person.get("role", ""),
            str(person["signed_days"]) + " 天",
            str(person["missing_days"]) + " 天",
            str(round(person["attendance_rate"] * 100)) + "%",
        ]
        status = person["status"]
        x = margin
        for i, (val, w) in enumerate(zip(values, col_widths)):
            color = (0x17, 0x25, 0x1F) if i == 0 else (0x4E, 0x60, 0x59)
            font = font_row_bold if i == 0 else font_row
            draw.text((x + 12, y + 8), val, fill=color, font=font)
            x += w
        # status badge
        sx = margin + sum(col_widths[:5]) + 12
        badge_colors = {
            "全勤": ((0xE8, 0xF5, 0xF0), (0x0D, 0x7A, 0x5C)),
            "部分签到": ((0xFF, 0xF7, 0xE8), (0xAD, 0x74, 0x18)),
            "未签到": ((0xFC, 0xE8, 0xE8), (0xD6, 0x30, 0x31)),
        }
        bc, tc = badge_colors.get(status, ((0xF0, 0xF4, 0xF2), (0x52, 0x61, 0x5C)))
        draw.rectangle([sx - 4, y + 6, sx + 60, y + row_h - 6], fill=bc)
        draw.text((sx, y + 8), status, fill=tc, font=font_row)
    # footer
    fy = header_h + row_h + len(people) * row_h + 14
    now_str = datetime.now().strftime("%Y-%m-%d %H:%M")
    draw.text((margin, fy), f"生成时间：{now_str}", fill=(0x71, 0x81, 0x7B), font=font_sub)
    buf = io.BytesIO()
    img.save(buf, "PNG")
    buf.seek(0)
    role_label = title_role if title_role != "全部" else "全部"
    filename = f"分包签到_{role_label}_{data['start']}_{data['end']}.png"
    return send_file(buf, as_attachment=True, download_name=filename, mimetype="image/png")


@app.get("/api/meta")
def meta():
    with db() as conn:
        categories = [r[0] for r in conn.execute("SELECT DISTINCT category FROM people WHERE active=1 ORDER BY category")]
        departments = [r[0] for r in conn.execute("SELECT DISTINCT department FROM people WHERE department<>'' ORDER BY department")]
        last_import = conn.execute("SELECT * FROM imports ORDER BY id DESC LIMIT 1").fetchone()
        bounds = conn.execute("SELECT MIN(check_date),MAX(check_date),COUNT(*) FROM hazards").fetchone()
    categories = sorted(set(categories) | set(ROLE_STANDARDS.keys()))
    return jsonify(categories=categories, departments=departments, defaultDepartments=DEFAULT_DEPARTMENTS,
                   datasetTypes=[{k: v for k, v in item.items() if k != "filenamePattern"} for item in DATASET_TYPES],
                   lastImport=dict(last_import) if last_import else None,
                   bounds={"min": bounds[0], "max": bounds[1], "count": bounds[2]})


@app.post("/api/import")
def upload():
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify(error="请选择 Excel 文件"), 400
    if not file.filename.lower().endswith(".xlsx"):
        return jsonify(error="目前只支持 .xlsx 文件"), 400
    dataset = detect_dataset_type(file.filename)
    if not dataset:
        examples = "、".join(item["filenameExample"] for item in DATASET_TYPES)
        return jsonify(error=f"无法根据文件名识别数据类型。当前支持：{examples}", code="unknown_dataset"), 400
    target = UPLOAD_DIR / f"{datetime.now():%Y%m%d%H%M%S%f}_{dataset['id']}.xlsx"
    file.save(target)
    try:
        if dataset["id"] == "hazard_entry":
            return jsonify(import_hazards(target, file.filename))
        return jsonify(error=f"数据类型“{dataset['label']}”尚未配置导入器"), 501
    except Exception as exc:
        target.unlink(missing_ok=True)
        return jsonify(error=str(exc)), 400


@app.get("/api/hazards")
def hazards():
    start, end = range_args()
    page = max(1, int(request.args.get("page", 1)))
    size = min(200, max(10, int(request.args.get("size", 50))))
    search = request.args.get("search", "").strip()
    unit = request.args.get("unit", "").strip()
    where = ["check_date BETWEEN ? AND ?"]
    params = [start, end]
    if search:
        where.append("(checker_name LIKE ? OR description LIKE ? OR hazard_no LIKE ? OR area LIKE ?)")
        params += [f"%{search}%"] * 4
    if unit:
        where.append("check_unit=?")
        params.append(unit)
    clause = " AND ".join(where)
    with db() as conn:
        total = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause}", params).fetchone()[0]
        rows = conn.execute(f"""SELECT hazard_no,check_date,checker_name,check_unit,check_department,hazard_level,
                             hazard_category,description,area,status,responsible_department,responsible_team
                             FROM hazards WHERE {clause} ORDER BY check_date DESC,hazard_no DESC LIMIT ? OFFSET ?""",
                            params + [size, (page - 1) * size]).fetchall()
        units = [r[0] for r in conn.execute("SELECT DISTINCT check_unit FROM hazards WHERE check_unit<>'' ORDER BY check_unit")]
    return jsonify(items=[dict(r) for r in rows], total=total, page=page, size=size, units=units)


@app.get("/api/hazards/stats")
def hazard_stats():
    start, end = range_args()
    search = request.args.get("search", "").strip()
    unit = request.args.get("unit", "").strip()
    where = ["check_date BETWEEN ? AND ?", NON_D_HAZARD_SQL]
    params = [start, end]
    if search:
        where.append("(checker_name LIKE ? OR description LIKE ? OR hazard_no LIKE ? OR area LIKE ?)")
        params += [f"%{search}%"] * 4
    if unit:
        where.append("check_unit=?")
        params.append(unit)
    clause = " AND ".join(where)
    with db() as conn:
        total = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause}", params).fetchone()[0]
        internal_count = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause} AND check_unit=?", params + ["中建二局"]).fetchone()[0]
        external_count = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause} AND check_unit=?", params + ["工程公司"]).fetchone()[0]
        b_hazards = [dict(r) for r in conn.execute(
            f"SELECT description, checker_name, check_date FROM hazards WHERE {clause} AND hazard_level='B' ORDER BY check_date DESC", params)]
        b_internal = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause} AND hazard_level='B' AND check_unit=?", params + ["中建二局"]).fetchone()[0]
        b_external = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause} AND hazard_level='B' AND check_unit=?", params + ["工程公司"]).fetchone()[0]
        rectification_count = conn.execute(f"SELECT COUNT(*) FROM hazards WHERE {clause} AND status=?", params + ["进行中"]).fetchone()[0]
        rectification_rate = round((1 - rectification_count / total) * 100, 1) if total > 0 else 0
    return jsonify(total=total, internal=internal_count, external=external_count, bHazards=b_hazards, bInternal=b_internal, bExternal=b_external, rectificationCount=rectification_count, rectificationRate=rectification_rate)


@app.get("/api/hazards/category-stats")
def hazard_category_stats():
    start, end = range_args()
    search = request.args.get("search", "").strip()
    unit = request.args.get("unit", "").strip()
    where = ["check_date BETWEEN ? AND ?", NON_D_HAZARD_SQL]
    params = [start, end]
    if search:
        where.append("(checker_name LIKE ? OR description LIKE ? OR hazard_no LIKE ? OR area LIKE ?)")
        params += [f"%{search}%"] * 4
    if unit:
        where.append("check_unit=?")
        params.append(unit)
    clause = " AND ".join(where)
    with db() as conn:
        cats = conn.execute(
            f"SELECT hazard_category, check_unit, COUNT(*) as cnt FROM hazards WHERE {clause} AND hazard_category IS NOT NULL AND hazard_category != '' GROUP BY hazard_category, check_unit ORDER BY cnt DESC",
            params).fetchall()
        levels = conn.execute(
            f"SELECT hazard_level, COUNT(*) as cnt FROM hazards WHERE {clause} AND hazard_level IS NOT NULL AND hazard_level != '' GROUP BY hazard_level ORDER BY cnt DESC",
            params).fetchall()
    cat_map = {}
    for row in cats:
        cat = row[0]
        unit = row[1]
        cnt = row[2]
        if cat not in cat_map:
            cat_map[cat] = {"category": cat, "total": 0, "internal": 0, "external": 0}
        cat_map[cat]["total"] += cnt
        if unit == "中建二局":
            cat_map[cat]["internal"] += cnt
        else:
            cat_map[cat]["external"] += cnt
    categories = sorted(cat_map.values(), key=lambda x: x["total"], reverse=True)
    levels_result = [{"level": row[0], "count": row[1]} for row in levels]
    return jsonify(categories=categories, levels=levels_result)


@app.get("/api/hazards/category-descriptions")
def hazard_category_descriptions():
    start, end = range_args()
    category = request.args.get("category", "").strip()
    if not category:
        return jsonify(descriptions=[])
    search = request.args.get("search", "").strip()
    unit = request.args.get("unit", "").strip()
    where = ["check_date BETWEEN ? AND ?", NON_D_HAZARD_SQL, "hazard_category=?"]
    params = [start, end, category]
    if search:
        where.append("(checker_name LIKE ? OR description LIKE ? OR hazard_no LIKE ? OR area LIKE ?)")
        params += [f"%{search}%"] * 4
    if unit:
        where.append("check_unit=?")
        params.append(unit)
    clause = " AND ".join(where)
    with db() as conn:
        rows = conn.execute(
            f"SELECT description, check_date, check_unit, hazard_level FROM hazards WHERE {clause} ORDER BY check_date DESC LIMIT 100",
            params).fetchall()
    return jsonify(descriptions=[dict(r) for r in rows])


@app.get("/api/people")
def people():
    with db() as conn:
        rows = conn.execute("SELECT * FROM people ORDER BY category,department,name").fetchall()
    return jsonify(items=[dict(r) for r in rows])


@app.post("/api/people/import")
def people_import():
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify(error="请选择人员 Excel 文件"), 400
    if not file.filename.lower().endswith(".xlsx"):
        return jsonify(error="目前只支持 .xlsx 文件"), 400
    mode = request.form.get("mode", "append")
    if mode not in ("append", "overwrite"):
        return jsonify(error="导入模式只能是 append 或 overwrite"), 400
    target = UPLOAD_DIR / f"{datetime.now():%Y%m%d%H%M%S%f}_people.xlsx"
    file.save(target)
    try:
        return jsonify(import_people(target, file.filename, mode))
    except Exception as exc:
        target.unlink(missing_ok=True)
        return jsonify(error=str(exc)), 400


@app.post("/api/people")
def save_person():
    item = request.get_json(force=True)
    name = clean_name(item.get("name"))
    if not name:
        return jsonify(error="姓名不能为空"), 400
    values = (name, str(item.get("category") or "其他"), str(item.get("department") or ""),
              max(0, float(item.get("target_count") or 0)),
              "day" if item.get("target_period") == "day" else "week", 1 if item.get("active", True) else 0)
    with db() as conn:
        if item.get("id"):
            conn.execute("UPDATE people SET name=?,category=?,department=?,target_count=?,target_period=?,active=? WHERE id=?",
                         values + (int(item["id"]),))
            person_id = int(item["id"])
        else:
            cursor = conn.execute("INSERT INTO people(name,category,department,target_count,target_period,active) VALUES (?,?,?,?,?,?)", values)
            person_id = cursor.lastrowid
    return jsonify(ok=True, id=person_id)


@app.delete("/api/people/<int:person_id>")
def delete_person(person_id):
    with db() as conn:
        conn.execute("DELETE FROM people WHERE id=?", (person_id,))
    return jsonify(ok=True)


@app.get("/api/statistics")
def statistics():
    start, end = range_args()
    category = request.args.get("category", "").strip()
    department = request.args.get("department", "").strip()
    start_date, end_date = date.fromisoformat(start), date.fromisoformat(end)
    days = (end_date - start_date).days + 1
    with db() as conn:
        settings = {r[0]: r[1] for r in conn.execute("SELECT key,value FROM settings")}
        internal = settings.get("internal_unit", "中建二局")
        ratio_target = float(settings.get("ratio_target", 5))
        internal_count = conn.execute(
            f"SELECT COUNT(*) FROM hazards WHERE check_date BETWEEN ? AND ? AND {NON_D_HAZARD_SQL} AND check_unit=?",
            (start, end, internal)).fetchone()[0]
        external_count = conn.execute(
            f"SELECT COUNT(*) FROM hazards WHERE check_date BETWEEN ? AND ? AND {NON_D_HAZARD_SQL} AND check_unit=?",
            (start, end, "工程公司")).fetchone()[0]
        where = ["active=1"]
        params = []
        if category:
            where.append("category=?"); params.append(category)
        if department:
            where.append("department=?"); params.append(department)
        roster = conn.execute(f"SELECT * FROM people WHERE {' AND '.join(where)} ORDER BY category,department,name", params).fetchall()
        counts = {r[0]: r[1] for r in conn.execute(
            f"SELECT checker_name,COUNT(*) FROM hazards WHERE check_date BETWEEN ? AND ? AND {NON_D_HAZARD_SQL} GROUP BY checker_name",
            (start, end))}
        b_counts = {r[0]: r[1] for r in conn.execute("SELECT checker_name,COUNT(*) FROM hazards WHERE check_date BETWEEN ? AND ? AND hazard_level='B' GROUP BY checker_name", (start, end))}
    result = []
    for person in roster:
        multiplier = days if person["target_period"] == "day" else math.ceil(days / 7)
        target = person["target_count"] * multiplier
        count = counts.get(person["name"], 0)
        result.append({**dict(person), "count": count, "bCount": b_counts.get(person["name"], 0),
                       "periodTarget": target, "met": count >= target})
    ratio = None if external_count == 0 else round(internal_count / external_count, 2)
    return jsonify(start=start, end=end, days=days, people=result,
                   comparison={"internal": internal_count, "external": external_count, "ratio": ratio,
                               "target": ratio_target, "met": external_count == 0 or ratio >= ratio_target,
                               "internalUnit": internal})


@app.get("/api/training/schedule")
def training_schedule():
    events = read_training_schedule()
    bounds = {
        "min": min((event["date"] for event in events), default=""),
        "max": max((event["date"] for event in events), default=""),
        "count": len(events),
        "source": TRAINING_SCHEDULE_FILE.name if TRAINING_SCHEDULE_FILE.exists() else "",
    }
    start = request.args.get("start") or ""
    end = request.args.get("end") or ""
    if not start or not end:
        today = date.today().isoformat()
        if bounds["min"] and not (bounds["min"] <= today <= bounds["max"]):
            today = bounds["min"]
        start, end = week_range_for(today)
    keyword = (request.args.get("keyword") or "").strip()
    filtered = []
    for event in events:
        if start and event["date"] < start:
            continue
        if end and event["date"] > end:
            continue
        haystack = " ".join([event["date"], event["weekday"], event["time"], event["period"], event["title"], " ".join(event["items"])])
        if keyword and keyword not in haystack:
            continue
        filtered.append(event)
    return jsonify({"items": filtered, "bounds": bounds, "start": start, "end": end})


@app.post("/api/training/save")
def training_save():
    body = request.get_json(force=True)
    pwd = str(body.get("password") or "")
    if pwd != TRAINING_EDIT_PASSWORD:
        return jsonify(error="密码错误"), 403
    event_id = str(body.get("id") or "").strip()
    if not event_id:
        return jsonify(error="缺少事件 ID"), 400
    if event_id.startswith("ledger:"):
        ledger_id = event_id.split(":", 1)[1]
        updates = []
        params = []
        field_map = {"title": "name", "time": "schedule_time", "period": "schedule_period"}
        for source, target in field_map.items():
            if source in body and body[source] is not None:
                updates.append(f"{target}=?")
                params.append(str(body[source]).strip())
        if updates:
            params.append(ledger_id)
            with db() as conn:
                conn.execute(f"UPDATE training_ledger_events SET {','.join(updates)} WHERE id=?", params)
        return jsonify(ok=True, id=event_id)
    overrides = load_training_overrides()
    entry = overrides.get(event_id, {})
    for key in ("time", "period", "title", "items"):
        if key in body and body[key] is not None:
            entry[key] = body[key]
    overrides[event_id] = entry
    save_training_overrides(overrides)
    return jsonify(ok=True, id=event_id)


@app.post("/api/training/reset")
def training_reset():
    body = request.get_json(force=True)
    pwd = str(body.get("password") or "")
    if pwd != TRAINING_EDIT_PASSWORD:
        return jsonify(error="密码错误"), 403
    event_id = str(body.get("id") or "").strip()
    overrides = load_training_overrides()
    if event_id:
        overrides.pop(event_id, None)
    else:
        overrides.clear()
    save_training_overrides(overrides)
    return jsonify(ok=True)


@app.post("/api/training/verify")
def training_verify():
    body = request.get_json(force=True)
    if str(body.get("password") or "") == TRAINING_EDIT_PASSWORD:
        return jsonify(ok=True)
    return jsonify(error="密码错误"), 403


def ledger_password_ok():
    password = request.form.get("password") if request.form else ""
    if not password and request.is_json:
        password = (request.get_json(silent=True) or {}).get("password", "")
    return str(password or "") == TRAINING_EDIT_PASSWORD


def ledger_file_dict(row):
    item = dict(row)
    item["download_url"] = f"/api/training-ledger/files/{item['id']}/download"
    item["preview_url"] = f"/api/training-ledger/files/{item['id']}/preview"
    return item


@app.get("/api/training-ledger/events")
def training_ledger_events():
    keyword = (request.args.get("keyword") or "").strip()
    category = (request.args.get("category") or "").strip()
    start_date = (request.args.get("start_date") or "").strip()
    end_date = (request.args.get("end_date") or "").strip()
    where = ""
    params = []
    if category:
        where = "WHERE e.category = ?"
        params.append(category)
    if start_date:
        prefix = "AND " if where else "WHERE "
        where += f"{prefix}e.training_date >= ?"
        params.append(start_date)
    if end_date:
        prefix = "AND " if where else "WHERE "
        where += f"{prefix}e.training_date <= ?"
        params.append(end_date)
    if keyword:
        terms = [term for term in re.split(r"\s+", keyword) if term]
        searchable = """
            (e.name LIKE ? OR e.training_date LIKE ? OR e.description LIKE ?
             OR e.training_location LIKE ? OR e.instructor LIKE ? OR e.audience LIKE ?
             OR CAST(e.participant_count AS TEXT) LIKE ?
             OR e.schedule_time LIKE ? OR e.schedule_period LIKE ?
             OR EXISTS (
                 SELECT 1 FROM training_ledger_files sf
                 WHERE sf.event_id=e.id
                   AND (sf.original_name LIKE ? OR sf.display_name LIKE ?)
             ))
        """
        prefix = "AND " if where else "WHERE "
        where += prefix + " AND ".join(searchable for _ in terms)
        for term in terms:
            params.extend([f"%{term}%"] * 11)
    with db() as conn:
        events = conn.execute(f"""
            SELECT e.*, COUNT(f.id) AS file_count
            FROM training_ledger_events e
            LEFT JOIN training_ledger_files f ON f.event_id=e.id
            {where}
            GROUP BY e.id
            ORDER BY e.training_date DESC, e.id DESC
        """, params).fetchall()
        result = []
        for event in events:
            files = conn.execute("""
                SELECT id,event_id,original_name,display_name,kind,content_type,size,created_at
                FROM training_ledger_files WHERE event_id=? ORDER BY id DESC
            """, (event["id"],)).fetchall()
            entry = dict(event)
            entry["files"] = [ledger_file_dict(row) for row in files]
            result.append(entry)
    return jsonify(items=result)


@app.post("/api/training-ledger/export")
def training_ledger_export():
    body = request.get_json(force=True)
    raw_ids = body.get("ids") or []
    try:
        event_ids = [int(value) for value in raw_ids]
    except (TypeError, ValueError):
        return jsonify(error="培训记录选择无效"), 400
    if not event_ids:
        return jsonify(error="请至少选择一项培训记录"), 400
    placeholders = ",".join("?" for _ in event_ids)
    with db() as conn:
        rows = conn.execute(f"""
            SELECT id,name,audience,training_location,instructor,participant_count,training_date
            FROM training_ledger_events
            WHERE id IN ({placeholders})
            ORDER BY training_date,id
        """, event_ids).fetchall()
    if not rows:
        return jsonify(error="未找到所选培训记录"), 404
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "培训台账"
    sheet.sheet_view.showGridLines = False
    sheet.merge_cells("A1:F1")
    sheet["A1"] = "培训台账"
    sheet["A1"].font = Font(name="微软雅黑", size=20, bold=True, color="FFFFFF")
    sheet["A1"].fill = PatternFill("solid", fgColor="176B57")
    sheet["A1"].alignment = Alignment(horizontal="center", vertical="center")
    sheet.row_dimensions[1].height = 38
    sheet.merge_cells("A2:F2")
    sheet["A2"] = f"共 {len(rows)} 场培训 · 生成时间：{datetime.now():%Y-%m-%d %H:%M}"
    sheet["A2"].font = Font(name="微软雅黑", size=10, color="52645E")
    sheet["A2"].fill = PatternFill("solid", fgColor="E9F5F0")
    sheet["A2"].alignment = Alignment(horizontal="center", vertical="center")
    headers = ["培训日期", "培训主题", "培训对象", "培训地点", "培训讲师", "培训人数"]
    sheet.append(headers)
    for row in rows:
        sheet.append([
            date.fromisoformat(row["training_date"]),
            row["name"],
            row["audience"] or "",
            row["training_location"] or "",
            row["instructor"] or "",
            int(row["participant_count"] or 0),
        ])
    header_fill = PatternFill("solid", fgColor="267B67")
    light_fill = PatternFill("solid", fgColor="F3F8F6")
    thin = Side(style="thin", color="D9E6E1")
    for cell in sheet[3]:
        cell.font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    for row_index in range(4, sheet.max_row + 1):
        if row_index % 2 == 0:
            for cell in sheet[row_index]:
                cell.fill = light_fill
        for cell in sheet[row_index]:
            cell.font = Font(name="微软雅黑", size=10, color="24332E")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(bottom=thin)
        sheet.cell(row_index, 1).number_format = "yyyy-mm-dd"
        sheet.row_dimensions[row_index].height = 30
    widths = [15, 34, 28, 26, 18, 13]
    for index, width in enumerate(widths, 1):
        sheet.column_dimensions[get_column_letter(index)].width = width
    sheet.freeze_panes = "A4"
    sheet.auto_filter.ref = f"A3:F{sheet.max_row}"
    output = BytesIO()
    workbook.save(output)
    output.seek(0)
    filename = f"培训台账_{datetime.now():%Y%m%d}.xlsx"
    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/api/training-ledger/events")
def training_ledger_create_event():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    training_date = str(body.get("training_date") or "").strip()
    description = str(body.get("description") or "").strip()
    schedule_time = str(body.get("schedule_time") or "19:30-21:00").strip()
    schedule_period = str(body.get("schedule_period") or "晚上").strip()
    training_location = str(body.get("training_location") or "").strip()
    instructor = str(body.get("instructor") or "").strip()
    audience = str(body.get("audience") or "").strip()
    category = str(body.get("category") or "专题培训").strip()
    try:
        participant_count = max(0, int(body.get("participant_count") or 0))
    except (TypeError, ValueError):
        return jsonify(error="培训人数必须是数字"), 400
    if not name:
        return jsonify(error="请输入培训名目"), 400
    try:
        date.fromisoformat(training_date)
    except ValueError:
        return jsonify(error="请选择正确的培训日期"), 400
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        cursor = conn.execute("""
            INSERT INTO training_ledger_events
            (name,training_date,description,schedule_time,schedule_period,
             training_location,instructor,audience,participant_count,category,created_at)
            VALUES (?,?,?,?,?,?,?,?,?,?,?)
        """, (name, training_date, description, schedule_time, schedule_period,
              training_location, instructor, audience, participant_count, category, now))
    return jsonify(ok=True, id=cursor.lastrowid)


@app.patch("/api/training-ledger/events/<int:event_id>")
def training_ledger_update_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    fields = {}
    for key in ("name", "training_date", "description", "training_location", "instructor", "audience", "category"):
        if key in body:
            fields[key] = str(body.get(key) or "").strip()
    if "name" in fields and not fields["name"]:
        return jsonify(error="请输入培训名称"), 400
    if "training_date" in fields:
        try:
            date.fromisoformat(fields["training_date"])
        except ValueError:
            return jsonify(error="请选择正确的培训日期"), 400
    try:
        if "participant_count" in body:
            fields["participant_count"] = max(0, int(body.get("participant_count") or 0))
    except (TypeError, ValueError):
        return jsonify(error="培训人数必须是数字"), 400
    if not fields:
        return jsonify(ok=True, id=event_id)
    assignments = ",".join(f"{key}=?" for key in fields)
    with db() as conn:
        cursor = conn.execute(
            f"UPDATE training_ledger_events SET {assignments} WHERE id=?",
            [*fields.values(), event_id],
        )
        event = conn.execute("SELECT name,training_date FROM training_ledger_events WHERE id=?", (event_id,)).fetchone()
        if event:
            files = conn.execute(
                "SELECT id,original_name,kind FROM training_ledger_files WHERE event_id=?",
                (event_id,),
            ).fetchall()
            year, month, day = map(int, event["training_date"].split("-"))
            for item in files:
                suffix = Path(item["original_name"]).suffix.lower()
                file_label = "照片" if item["kind"] == "image" else "签到单"
                display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
                conn.execute(
                    "UPDATE training_ledger_files SET display_name=? WHERE id=?",
                    (display_name, item["id"]),
                )
    if not cursor.rowcount:
        return jsonify(error="培训名目不存在"), 404
    return jsonify(ok=True, id=event_id)


@app.delete("/api/training-ledger/events/<int:event_id>")
def training_ledger_delete_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        event = conn.execute("SELECT id FROM training_ledger_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="培训名目不存在"), 404
        stored_names = [
            row["stored_name"]
            for row in conn.execute("SELECT stored_name FROM training_ledger_files WHERE event_id=?", (event_id,))
        ]
        conn.execute("DELETE FROM training_ledger_events WHERE id=?", (event_id,))
    for stored_name in stored_names:
        target = TRAINING_LEDGER_DIR / stored_name
        if target.parent == TRAINING_LEDGER_DIR and target.is_file():
            target.unlink()
    return jsonify(ok=True, id=event_id)


@app.get("/api/training-ledger/categories")
def training_ledger_categories():
    with db() as conn:
        rows = conn.execute(
            "SELECT id, name, sort_order FROM training_categories ORDER BY sort_order, id"
        ).fetchall()
    return jsonify(items=[dict(row) for row in rows])


@app.post("/api/training-ledger/categories")
def training_ledger_create_category():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        existing = conn.execute(
            "SELECT id FROM training_categories WHERE name=?", (name,)
        ).fetchone()
        if existing:
            return jsonify(error="模块名称已存在"), 409
        max_order = conn.execute(
            "SELECT COALESCE(MAX(sort_order), -1) AS max_order FROM training_categories"
        ).fetchone()["max_order"]
        cursor = conn.execute(
            "INSERT INTO training_categories (name, sort_order, created_at) VALUES (?, ?, ?)",
            (name, max_order + 1, now),
        )
    return jsonify(ok=True, id=cursor.lastrowid, name=name, sort_order=max_order + 1)


@app.patch("/api/training-ledger/categories/<int:category_id>")
def training_ledger_update_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    with db() as conn:
        cat = conn.execute(
            "SELECT id, name FROM training_categories WHERE id=?", (category_id,)
        ).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        dup = conn.execute(
            "SELECT id FROM training_categories WHERE name=? AND id!=?", (name, category_id)
        ).fetchone()
        if dup:
            return jsonify(error="模块名称已存在"), 409
        conn.execute(
            "UPDATE training_categories SET name=? WHERE id=?", (name, category_id)
        )
        conn.execute(
            "UPDATE training_ledger_events SET category=? WHERE category=?",
            (name, cat["name"]),
        )
    return jsonify(ok=True, id=category_id, name=name)


@app.delete("/api/training-ledger/categories/<int:category_id>")
def training_ledger_delete_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        cat = conn.execute(
            "SELECT id, name FROM training_categories WHERE id=?", (category_id,)
        ).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        count = conn.execute("SELECT COUNT(*) AS cnt FROM training_categories").fetchone()["cnt"]
        if count <= 1:
            return jsonify(error="至少保留一个培训模块"), 400
        conn.execute(
            "UPDATE training_ledger_events SET category='专题培训' WHERE category=?",
            (cat["name"],)
        )
        conn.execute("DELETE FROM training_categories WHERE id=?", (category_id,))
    return jsonify(ok=True, id=category_id)


@app.post("/api/training-ledger/categories/reorder")
def training_ledger_reorder_categories():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    ids = body.get("ids")
    if not ids or not isinstance(ids, list):
        return jsonify(error="请提供模块ID列表"), 400
    with db() as conn:
        for index, cat_id in enumerate(ids):
            conn.execute(
                "UPDATE training_categories SET sort_order=? WHERE id=?",
                (index, int(cat_id)),
            )
    return jsonify(ok=True)


@app.get("/api/training-ledger/stats")
def training_ledger_stats():
    start_date = (request.args.get("start_date") or "").strip()
    end_date = (request.args.get("end_date") or "").strip()
    category = (request.args.get("category") or "").strip()
    where = []
    params = []
    if start_date:
        where.append("training_date >= ?")
        params.append(start_date)
    if end_date:
        where.append("training_date <= ?")
        params.append(end_date)
    if category:
        where.append("category = ?")
        params.append(category)
    clause = ("WHERE " + " AND ".join(where)) if where else ""
    with db() as conn:
        row = conn.execute(
            f"SELECT COUNT(*) AS sessions, COALESCE(SUM(participant_count), 0) AS participants FROM training_ledger_events {clause}",
            params,
        ).fetchone()
    return jsonify(sessions=row["sessions"], participants=row["participants"])


# ── Brake Warning Ledger ──────────────────────────────────────────────

@app.get("/brake-ledger")
def brake_ledger_page():
    return send_from_directory("public", "brake-ledger.html")


@app.get("/api/brake-ledger/categories")
def brake_ledger_categories():
    with db() as conn:
        rows = conn.execute(
            "SELECT id, name, sort_order FROM brake_ledger_categories ORDER BY sort_order, id"
        ).fetchall()
    return jsonify(items=[dict(row) for row in rows])


@app.post("/api/brake-ledger/categories")
def brake_ledger_create_category():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        if conn.execute("SELECT id FROM brake_ledger_categories WHERE name=?", (name,)).fetchone():
            return jsonify(error="模块名称已存在"), 409
        max_order = conn.execute("SELECT COALESCE(MAX(sort_order), -1) AS o FROM brake_ledger_categories").fetchone()["o"]
        cursor = conn.execute(
            "INSERT INTO brake_ledger_categories (name, sort_order, created_at) VALUES (?, ?, ?)",
            (name, max_order + 1, now),
        )
    return jsonify(ok=True, id=cursor.lastrowid, name=name, sort_order=max_order + 1)


@app.patch("/api/brake-ledger/categories/<int:category_id>")
def brake_ledger_update_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    with db() as conn:
        cat = conn.execute("SELECT id, name FROM brake_ledger_categories WHERE id=?", (category_id,)).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        if conn.execute("SELECT id FROM brake_ledger_categories WHERE name=? AND id!=?", (name, category_id)).fetchone():
            return jsonify(error="模块名称已存在"), 409
        conn.execute("UPDATE brake_ledger_categories SET name=? WHERE id=?", (name, category_id))
        conn.execute("UPDATE brake_ledger_events SET category=? WHERE category=?", (name, cat["name"]))
    return jsonify(ok=True, id=category_id, name=name)


@app.delete("/api/brake-ledger/categories/<int:category_id>")
def brake_ledger_delete_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        cat = conn.execute("SELECT id, name FROM brake_ledger_categories WHERE id=?", (category_id,)).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        if conn.execute("SELECT COUNT(*) AS cnt FROM brake_ledger_categories").fetchone()["cnt"] <= 1:
            return jsonify(error="至少保留一个模块"), 400
        conn.execute("UPDATE brake_ledger_events SET category=? WHERE category=?", (brake_categories[0][0], cat["name"]))
        conn.execute("DELETE FROM brake_ledger_categories WHERE id=?", (category_id,))
    return jsonify(ok=True, id=category_id)


@app.post("/api/brake-ledger/categories/reorder")
def brake_ledger_reorder_categories():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    ids = body.get("ids")
    if not ids or not isinstance(ids, list):
        return jsonify(error="请提供模块ID列表"), 400
    with db() as conn:
        for index, cid in enumerate(ids):
            conn.execute("UPDATE brake_ledger_categories SET sort_order=? WHERE id=?", (index, int(cid)))
    return jsonify(ok=True)


def _brake_file_dict(row):
    item = dict(row)
    item["download_url"] = f"/api/brake-ledger/files/{item['id']}/download"
    item["preview_url"] = f"/api/brake-ledger/files/{item['id']}/preview"
    return item


@app.get("/api/brake-ledger/events")
def brake_ledger_events():
    keyword = (request.args.get("keyword") or "").strip()
    category = (request.args.get("category") or "").strip()
    where = ""
    params = []
    if category:
        where = "WHERE e.category = ?"
        params.append(category)
    if keyword:
        terms = [t for t in re.split(r"\s+", keyword) if t]
        searchable = """
            (e.name LIKE ? OR e.record_date LIKE ? OR e.description LIKE ?
             OR e.issue_dept LIKE ? OR e.responsible_dept LIKE ? OR e.responsible_person LIKE ?
             OR e.area LIKE ? OR e.subcontractor LIKE ? OR e.team LIKE ?
             OR EXISTS (
                 SELECT 1 FROM brake_ledger_files sf
                 WHERE sf.event_id=e.id
                   AND (sf.original_name LIKE ? OR sf.display_name LIKE ?)
             ))
        """
        prefix = "AND " if where else "WHERE "
        where += prefix + " AND ".join(searchable for _ in terms)
        for term in terms:
            params.extend([f"%{term}%"] * 11)
    with db() as conn:
        events = conn.execute(f"""
            SELECT e.*, COUNT(f.id) AS file_count
            FROM brake_ledger_events e
            LEFT JOIN brake_ledger_files f ON f.event_id=e.id
            {where}
            GROUP BY e.id
            ORDER BY e.record_date DESC, e.id DESC
        """, params).fetchall()
        result = []
        for event in events:
            files = conn.execute(
                "SELECT id,event_id,original_name,display_name,kind,content_type,size,created_at FROM brake_ledger_files WHERE event_id=? ORDER BY id DESC",
                (event["id"],),
            ).fetchall()
            entry = dict(event)
            entry["files"] = [_brake_file_dict(row) for row in files]
            result.append(entry)
    return jsonify(items=result)


@app.post("/api/brake-ledger/events")
def brake_ledger_create_event():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    record_date = str(body.get("record_date") or "").strip()
    if not name:
        return jsonify(error="请输入记录名称"), 400
    try:
        date.fromisoformat(record_date)
    except ValueError:
        return jsonify(error="请选择正确的日期"), 400
    category = str(body.get("category") or "工程公司整改单").strip()
    description = str(body.get("description") or "").strip()
    issue_dept = str(body.get("issue_dept") or "").strip()
    responsible_dept = str(body.get("responsible_dept") or "").strip()
    responsible_person = str(body.get("responsible_person") or "").strip()
    area = str(body.get("area") or "").strip()
    subcontractor = str(body.get("subcontractor") or "").strip()
    team = str(body.get("team") or "").strip()
    status = str(body.get("status") or "").strip()
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        cursor = conn.execute(
            "INSERT INTO brake_ledger_events (name,record_date,category,description,issue_dept,responsible_dept,responsible_person,area,subcontractor,team,status,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?)",
            (name, record_date, category, description, issue_dept, responsible_dept, responsible_person, area, subcontractor, team, status, now),
        )
    return jsonify(ok=True, id=cursor.lastrowid)


@app.patch("/api/brake-ledger/events/<int:event_id>")
def brake_ledger_update_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    fields = {}
    for key in ("name", "record_date", "category", "description", "issue_dept", "responsible_dept", "responsible_person", "area", "subcontractor", "team", "status"):
        if key in body:
            fields[key] = str(body.get(key) or "").strip()
    if "name" in fields and not fields["name"]:
        return jsonify(error="请输入记录名称"), 400
    if "record_date" in fields:
        try:
            date.fromisoformat(fields["record_date"])
        except ValueError:
            return jsonify(error="请选择正确的日期"), 400
    if not fields:
        return jsonify(ok=True, id=event_id)
    assignments = ",".join(f"{k}=?" for k in fields)
    with db() as conn:
        cursor = conn.execute(
            f"UPDATE brake_ledger_events SET {assignments} WHERE id=?",
            [*fields.values(), event_id],
        )
        event = conn.execute("SELECT name,record_date FROM brake_ledger_events WHERE id=?", (event_id,)).fetchone()
        if event:
            files = conn.execute("SELECT id,original_name,kind FROM brake_ledger_files WHERE event_id=?", (event_id,)).fetchall()
            year, month, day = map(int, event["record_date"].split("-"))
            for item in files:
                suffix = Path(item["original_name"]).suffix.lower()
                file_label = "照片" if item["kind"] == "image" else "文件"
                display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
                conn.execute("UPDATE brake_ledger_files SET display_name=? WHERE id=?", (display_name, item["id"]))
    if not cursor.rowcount:
        return jsonify(error="记录不存在"), 404
    return jsonify(ok=True, id=event_id)


@app.delete("/api/brake-ledger/events/<int:event_id>")
def brake_ledger_delete_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        event = conn.execute("SELECT id FROM brake_ledger_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="记录不存在"), 404
        stored_names = [row["stored_name"] for row in conn.execute("SELECT stored_name FROM brake_ledger_files WHERE event_id=?", (event_id,))]
        conn.execute("DELETE FROM brake_ledger_events WHERE id=?", (event_id,))
    for name in stored_names:
        target = BRAKE_LEDGER_DIR / name
        if target.parent == BRAKE_LEDGER_DIR and target.is_file():
            target.unlink()
    return jsonify(ok=True, id=event_id)


@app.post("/api/brake-ledger/events/<int:event_id>/files")
def brake_ledger_upload(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    files = [f for f in request.files.getlist("files") if f and f.filename]
    if not files:
        return jsonify(error="请选择要上传的文件"), 400
    with db() as conn:
        event = conn.execute("SELECT * FROM brake_ledger_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="记录不存在"), 404
        saved = []
        for upload_file in files:
            original = Path(upload_file.filename).name
            safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original).strip(" .") or "文件"
            suffix = Path(safe_original).suffix.lower()
            stored_name = f"{event_id}_{uuid.uuid4().hex}{suffix}"
            target = BRAKE_LEDGER_DIR / stored_name
            upload_file.save(target)
            ct = upload_file.mimetype or "application/octet-stream"
            if ct.startswith("image/"):
                kind = "image"
            elif ct == "application/pdf":
                kind = "pdf"
            elif ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
                kind = "docx"
            else:
                kind = "file"
            year, month, day = map(int, event["record_date"].split("-"))
            file_label = "照片" if kind == "image" else "文件"
            display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
            now = datetime.now().isoformat(timespec="seconds")
            cursor = conn.execute(
                "INSERT INTO brake_ledger_files (event_id,original_name,stored_name,display_name,kind,content_type,size,created_at) VALUES (?,?,?,?,?,?,?,?)",
                (event_id, original, stored_name, display_name, kind, ct, upload_file.tell(), now),
            )
            saved.append({"id": cursor.lastrowid, "display_name": display_name})
    return jsonify(ok=True, items=saved)


@app.get("/api/brake-ledger/files/<int:file_id>/download")
def brake_ledger_download(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM brake_ledger_files WHERE id=?", (file_id,)).fetchone()
    if not item:
        return jsonify(error="文件不存在"), 404
    target = BRAKE_LEDGER_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    return send_file(target, as_attachment=True, download_name=item["display_name"], mimetype=item["content_type"])


@app.get("/api/brake-ledger/files/<int:file_id>/preview")
def brake_ledger_preview(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM brake_ledger_files WHERE id=?", (file_id,)).fetchone()
    if not item:
        return jsonify(error="文件不存在"), 404
    target = BRAKE_LEDGER_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    if item["kind"] == "image":
        return send_file(target, mimetype=item["content_type"], conditional=True)
    if item["kind"] == "pdf":
        return send_file(target, mimetype="application/pdf", conditional=True)
    if item["kind"] == "docx":
        try:
            from docx import Document
            doc = Document(target)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return jsonify(text=text, kind="docx")
        except Exception as e:
            return jsonify(error=f"DOCX 解析失败: {str(e)}"), 500
    return jsonify(error="不支持预览此文件类型"), 400


@app.delete("/api/brake-ledger/files/<int:file_id>")
def brake_ledger_delete_file(file_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        item = conn.execute("SELECT * FROM brake_ledger_files WHERE id=?", (file_id,)).fetchone()
        if not item:
            return jsonify(error="文件不存在"), 404
        conn.execute("DELETE FROM brake_ledger_files WHERE id=?", (file_id,))
    target = BRAKE_LEDGER_DIR / item["stored_name"]
    if target.is_file():
        target.unlink()
    return jsonify(ok=True, id=file_id)


@app.get("/api/brake-ledger/stats")
def brake_ledger_stats():
    start_date = (request.args.get("start_date") or "").strip()
    end_date = (request.args.get("end_date") or "").strip()
    category = (request.args.get("category") or "").strip()
    where = []
    params = []
    if start_date:
        where.append("record_date >= ?")
        params.append(start_date)
    if end_date:
        where.append("record_date <= ?")
        params.append(end_date)
    if category:
        where.append("category = ?")
        params.append(category)
    clause = ("WHERE " + " AND ".join(where)) if where else ""
    with db() as conn:
        row = conn.execute(
            f"SELECT COUNT(*) AS records, (SELECT COUNT(*) FROM brake_ledger_files{(' WHERE event_id IN (SELECT id FROM brake_ledger_events '+clause+')') if clause else ''}) AS files FROM brake_ledger_events {clause}",
            params,
        ).fetchone()
    return jsonify(records=row["records"], files=row["files"])


@app.post("/api/brake-ledger/clear")
def brake_ledger_clear():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        files = conn.execute("SELECT stored_name FROM brake_ledger_files").fetchall()
        conn.execute("DELETE FROM brake_ledger_events")
        conn.execute("DELETE FROM brake_ledger_files")
    for row in files:
        target = BRAKE_LEDGER_DIR / row["stored_name"]
        if target.parent == BRAKE_LEDGER_DIR and target.is_file():
            target.unlink()
    return jsonify(ok=True, deleted=len(files))


# ── Matching table for subcontractor lookup ──────────────────────────

_matching_cache = None  # {name: [{sub, team, group, team_code}], team_code: [{sub, ...}]}


MATCHING_TABLE_PATH = ROOT / "匹配台账.xlsx"


def _load_matching_table():
    """Parse 匹配台账.xlsx and return a lookup dict."""
    global _matching_cache
    if not MATCHING_TABLE_PATH.is_file():
        return None
    try:
        wb = load_workbook(MATCHING_TABLE_PATH, data_only=True)
        ws = wb.active
        rows = list(ws.iter_rows(min_row=4, values_only=True))
        wb.close()
    except Exception:
        return _matching_cache  # return stale cache if any

    by_name = {}
    by_team = {}
    for row in rows:
        name = str(row[7]).strip() if row[7] else ""  # col 8 (0-indexed: 7)
        if not name:
            continue
        group_code = str(row[4]).strip() if row[4] else ""  # col 5
        team_code = str(row[5]).strip() if row[5] else ""  # col 6
        sub = str(row[36]).strip() if row[36] else ""  # col 37: 分包名称（大班组名称）
        entry = {"sub": sub, "group": group_code, "team_code": team_code}

        by_name.setdefault(name, []).append(entry)
        if team_code:
            by_team.setdefault(team_code, []).append(entry)

    _matching_cache = {"by_name": by_name, "by_team": by_team}
    return _matching_cache


def _get_matching_table():
    global _matching_cache
    if _matching_cache is None:
        _load_matching_table()
    return _matching_cache


def match_subcontractor(person_name, team_name):
    """Look up subcontractor and team from matching table.
    Returns (subcontractor_str, team_str) — comma-separated if multiple matches.
    """
    table = _get_matching_table()
    if not table:
        return "", ""
    subs = set()
    teams = set()

    # Match by person name
    if person_name:
        entries = table.get("by_name", {}).get(person_name, [])
        for e in entries:
            if e["sub"]:
                subs.add(e["sub"])
            if e["team_code"]:
                teams.add(e["team_code"])

    # Match by team name (search in team_code and sub)
    if team_name:
        for e in table.get("by_team", {}).get(team_name, []):
            if e["sub"]:
                subs.add(e["sub"])
            if e["team_code"]:
                teams.add(e["team_code"])

    return ", ".join(sorted(subs)), ", ".join(sorted(teams))


@app.post("/api/brake-ledger/matching-table")
def brake_ledger_matching_table():
    """Upload/update the matching reference table."""
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    file = request.files.get("file")
    if not file:
        return jsonify(error="请选择匹配台账 Excel 文件"), 400
    try:
        file.save(str(MATCHING_TABLE_PATH))
        _matching_cache = None
        table = _load_matching_table()
        count = len(table.get("by_name", {})) if table else 0
        return jsonify(ok=True, people_count=count)
    except Exception as e:
        return jsonify(error=str(e)), 500


@app.post("/api/brake-ledger/import")
def brake_ledger_import():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    file = request.files.get("file")
    if not file:
        return jsonify(error="请选择一个 Excel 文件"), 400

    file_bytes = file.read()
    try:
        wb = load_workbook(BytesIO(file_bytes), data_only=True)
    except Exception:
        return jsonify(error="文件不是有效的 Excel 格式"), 400

    ws = wb.active
    rows = list(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 5000), values_only=True))
    if len(rows) < 2:
        return jsonify(error="Excel 至少需要标题行和一行数据"), 400

    raw_headers = [str(c).strip() if c else "" for c in rows[0]]
    # The actual headers are usually in row 2; row 1 is the sheet title
    if len(rows) >= 2:
        row2 = [str(c).strip() if c else "" for c in rows[1]]
        if any(h in " ".join(row2) for h in ["序号", "编号", "纠偏编号", "约谈编号"]):
            raw_headers = row2

    # Determine file type from all rows (title row + headers)
    headers = raw_headers
    header_str = " ".join(headers)
    # Also check the title row for sheet-level keywords
    title_str = " ".join(str(c).strip() if c else "" for c in rows[0])

    # ── Detect file type ──
    combined = header_str + " " + title_str
    if "偏差" in combined and "纠偏编号" in header_str:
        file_type = "deviation"
    elif "约谈编号" in header_str or ("约谈主题" in header_str and "被约谈方" in header_str):
        file_type = "interview"
    elif "停工" in combined and "复工条件" in header_str:
        file_type = "stop_work"
    elif "挂牌督办" in combined:
        file_type = "supervision"
    elif "通报批评" in combined:
        file_type = "criticism"
    elif "整改" in combined:
        file_type = "rectify"
    else:
        return jsonify(error="无法识别台账类型，请确认文件内容"), 400

    def col(name, default=""):
        """Find column index by header name (exact match first, then partial match)."""
        for i, h in enumerate(headers):
            if h == name:
                return i
        for i, h in enumerate(headers):
            if name in h:
                return i
        return -1

    def val(row, name, default=""):
        idx = col(name)
        if idx >= 0 and idx < len(row):
            v = row[idx]
            return str(v).strip() if v is not None else default
        return default

    overwrite = request.form.get("overwrite", "").strip().lower() == "true"
    imported = 0
    overwritten = 0
    skipped_existing = 0
    skipped_filter = 0
    file_attached = False

    # Pre-load matching table for deviation imports
    _get_matching_table()

    with db() as conn:
        for row in rows[1:]:
            if not row or all(c is None for c in row):
                continue

            if file_type == "deviation":
                # Filter: 责任单位 must be 中建二局
                responsible = val(row, "责任单位")
                if "中建二局" not in responsible:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "纠偏编号")
                treatment = val(row, "处理措施")
                if "红牌" in treatment or "黄牌" in treatment:
                    category = "红黄牌"
                else:
                    category = "行为偏差"

                name = val(row, "纠偏编号")
                record_date = val(row, "偏差时间")
                description = val(row, "偏差描述")
                issue_dept = val(row, "发出方")
                responsible_person = val(row, "责任人")
                responsible_dept = val(row, "责任部门")
                area = val(row, "发生地点")
                team = val(row, "责任班组")
                # Look up subcontractor and team from matching table
                subcontractor, matched_team = match_subcontractor(responsible_person, team)
                if matched_team:
                    team = matched_team

            elif file_type == "interview":
                # Filter: 被约谈方 contains 中建二局
                party = val(row, "被约谈方")
                if "中建二局" not in party:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "约谈编号")
                category = "约谈记录"
                name = val(row, "约谈主题")
                record_date = val(row, "约谈日期")
                description = val(row, "约谈纪要")
                issue_dept = val(row, "主持人")
                responsible_person = val(row, "被约谈人")
                responsible_dept = ""
                area = val(row, "约谈地点")
                subcontractor = val(row, "被约谈方")
                team = ""

            elif file_type == "stop_work":
                responsible = val(row, "责任单位")
                if "中建二局" not in responsible:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "编号")
                category = "工程公司停工令"
                name = val(row, "主题")
                record_date = val(row, "发出日期")
                description = val(row, "停工原因及依据")
                issue_dept = val(row, "发出方")
                responsible_person = val(row, "录入人")
                responsible_dept = ""
                area = val(row, "停工范围")
                subcontractor = val(row, "责任单位")
                team = ""

            elif file_type == "supervision":
                responsible = val(row, "责任承包商")
                if "中建二局" not in responsible:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "编号")
                category = "工程公司挂牌督办单"
                name = val(row, "主题")
                record_date = val(row, "发出日期")
                description = val(row, "问题描述")
                issue_dept = val(row, "发出方")
                responsible_person = val(row, "发出人")
                responsible_dept = val(row, "内部责任方")
                area = val(row, "涉及厂房")
                subcontractor = val(row, "责任承包商")
                team = ""

            elif file_type == "criticism":
                responsible = val(row, "责任承包商")
                if "中建二局" not in responsible:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "编号")
                category = "工程公司通报批评"
                name = val(row, "主题")
                record_date = val(row, "发出日期")
                description = val(row, "问题描述")
                issue_dept = val(row, "发出方")
                responsible_person = val(row, "发出人")
                responsible_dept = val(row, "内部责任方")
                area = val(row, "涉及厂房")
                subcontractor = val(row, "责任承包商")
                team = ""

            elif file_type == "rectify":
                responsible = val(row, "责任承包商")
                if "中建二局" not in responsible:
                    skipped_filter += 1
                    continue
                status = val(row, "状态")
                if status == "作废":
                    skipped_filter += 1
                    continue

                source_ref = val(row, "编号")
                category = "工程公司整改单"
                name = val(row, "主题")
                record_date = val(row, "发出日期")
                description = val(row, "问题描述")
                issue_dept = val(row, "发出方")
                responsible_person = val(row, "发出人")
                responsible_dept = val(row, "内部责任方")
                area = val(row, "涉及厂房")
                subcontractor = val(row, "责任承包商")
                team = ""

            else:
                continue

            if not name or not source_ref:
                skipped_filter += 1
                continue

            # Check for duplicate
            existing = conn.execute(
                "SELECT id FROM brake_ledger_events WHERE source_ref = ? AND source_ref != ''",
                (source_ref,),
            ).fetchone()

            # Parse date
            date_str = record_date[:10] if record_date else ""
            if len(date_str) < 10:
                date_str = ""
            now = datetime.now().isoformat(timespec="seconds")

            if existing:
                event_id = existing["id"]
                if overwrite:
                    conn.execute(
                        """UPDATE brake_ledger_events
                           SET name=?, record_date=?, category=?, description=?, issue_dept=?,
                               responsible_dept=?, responsible_person=?, area=?, subcontractor=?,
                               team=?, status=?, created_at=?
                           WHERE id=?""",
                        (name, date_str, category, description, issue_dept, responsible_dept,
                         responsible_person, area, subcontractor, team, status, now, event_id),
                    )
                    overwritten += 1
                else:
                    skipped_existing += 1
                # Attach the source xlsx file to the existing event
                if not file_attached:
                    try:
                        original_name = Path(file.filename).name if file.filename else "source.xlsx"
                        safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original_name).strip(" .") or "文件"
                        sfx = Path(safe_original).suffix.lower()
                        sname = f"{event_id}_{uuid.uuid4().hex}{sfx}"
                        target = BRAKE_LEDGER_DIR / sname
                        target.write_bytes(file_bytes)
                        try:
                            y, m, d = map(int, date_str.split("-"))
                        except (ValueError, AttributeError):
                            y, m, d = 2026, 1, 1
                        dname = f"{y}年{m}月{d}日{name}文件{sfx}"
                        conn.execute(
                            "INSERT INTO brake_ledger_files (event_id, original_name, stored_name, display_name, kind, content_type, size, created_at) VALUES (?,?,?,?,?,?,?,?)",
                            (event_id, original_name, sname, dname, "file", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", len(file_bytes), now),
                        )
                        file_attached = True
                    except Exception:
                        pass
                continue

            try:
                cursor = conn.execute(
                    """INSERT INTO brake_ledger_events
                       (name, record_date, category, description, issue_dept, responsible_dept,
                        responsible_person, area, subcontractor, team, status, source_ref, created_at)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)""",
                    (name, date_str, category, description, issue_dept, responsible_dept,
                     responsible_person, area, subcontractor, team, status, source_ref, now),
                )
                event_id = cursor.lastrowid
                imported += 1
                # Attach source file to new event
                if not file_attached:
                    try:
                        original_name = Path(file.filename).name if file.filename else "source.xlsx"
                        safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original_name).strip(" .") or "文件"
                        sfx = Path(safe_original).suffix.lower()
                        sname = f"{event_id}_{uuid.uuid4().hex}{sfx}"
                        target = BRAKE_LEDGER_DIR / sname
                        target.write_bytes(file_bytes)
                        try:
                            y, m, d = map(int, date_str.split("-"))
                        except (ValueError, AttributeError):
                            y, m, d = 2026, 1, 1
                        dname = f"{y}年{m}月{d}日{name}文件{sfx}"
                        conn.execute(
                            "INSERT INTO brake_ledger_files (event_id, original_name, stored_name, display_name, kind, content_type, size, created_at) VALUES (?,?,?,?,?,?,?,?)",
                            (event_id, original_name, sname, dname, "file", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", len(file_bytes), now),
                        )
                        file_attached = True
                    except Exception:
                        pass
            except Exception:
                skipped_existing += 1

    return jsonify(
        ok=True,
        imported=imported,
        overwritten=overwritten,
        skipped_existing=skipped_existing,
        skipped_filter=skipped_filter,
        file_attached=file_attached,
        category=category if (imported or overwritten) else "",
    )


@app.post("/api/brake-ledger/export")
def brake_ledger_export():
    body = request.get_json(force=True) or {}
    raw_ids = body.get("ids") or []
    try:
        event_ids = [int(v) for v in raw_ids]
    except (TypeError, ValueError):
        return jsonify(error="记录选择无效"), 400
    with db() as conn:
        if event_ids:
            placeholders = ",".join("?" for _ in event_ids)
            rows = conn.execute(
                f"SELECT * FROM brake_ledger_events WHERE id IN ({placeholders}) ORDER BY record_date, id",
                event_ids,
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT * FROM brake_ledger_events ORDER BY record_date, id"
            ).fetchall()
    if not rows:
        return jsonify(error="未找到记录"), 404

    wb = Workbook()
    wb.remove(wb.active)
    header_font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="267B67")
    cell_font = Font(name="微软雅黑", size=10, color="24332E")
    light_fill = PatternFill("solid", fgColor="F3F8F6")
    thin = Side(style="thin", color="D9E6E1")

    # Sheet definition: (sheet_name, [(header, event_field), ...])
    # event_field can be a DB column name, "seq" for auto-number, or "" for blank
    sheet_specs = [
        ("工程公司挂牌督办单", [
            ("序号", "seq"), ("签发部门", "issue_dept"), ("主题", "name"),
            ("发出日期", "record_date"), ("责任部门", "responsible_dept"),
            ("责任人", "responsible_person"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
            ("是否关闭", "status"), ("备注", "description"),
        ]),
        ("红黄牌", [
            ("序号", "seq"), ("发生地点", "area"), ("发出日期", "record_date"),
            ("责任人", "responsible_person"), ("责任部门", "responsible_dept"),
            ("责任班组", "team"), ("偏差类型", ""), ("偏差等级", ""),
            ("偏差描述", "description"), ("处理措施", ""), ("积分分值", ""),
            ("分包单位", "subcontractor"),
        ]),
        ("工程公司处理通报", [
            ("序号", "seq"), ("主题", "name"), ("发出日期", "record_date"),
            ("责任部门", "responsible_dept"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
            ("责任人", "responsible_person"), ("是否关闭", "status"),
            ("备注", "description"),
        ]),
        ("工程公司通报批评", [
            ("序号", "seq"), ("签发部门", "issue_dept"), ("主题", "name"),
            ("发出日期", "record_date"), ("责任部门", "responsible_dept"),
            ("责任人", "responsible_person"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
            ("是否关闭", "status"), ("备注", "description"),
        ]),
        ("工程公司整改单", [
            ("序号", "seq"), ("签发部门", "issue_dept"), ("主题", "name"),
            ("发出日期", "record_date"), ("责任部门", "responsible_dept"),
            ("责任人", "responsible_person"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
            ("是否关闭", "status"), ("备注", "description"),
        ]),
        ("工程公司停工令", [
            ("序号", "seq"), ("主题", "name"), ("发出部门", "issue_dept"),
            ("停工时间", "record_date"), ("停工原因", "description"),
            ("是否复工", "status"), ("责任部门", "responsible_dept"),
            ("责任人", "responsible_person"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
        ]),
        ("监理业主整改通知单", [
            ("序号", "seq"), ("签发部门", "issue_dept"), ("主题", "name"),
            ("发出日期", "record_date"), ("责任部门", "responsible_dept"),
            ("责任人", "responsible_person"), ("涉及区域", "area"),
            ("涉及分包", "subcontractor"), ("涉及班组", "team"),
            ("是否关闭", "status"), ("备注", "description"),
        ]),
        ("行为偏差", [
            ("序号", "seq"), ("发生地点", "area"), ("偏差时间", "record_date"),
            ("责任人", "responsible_person"), ("责任部门", "responsible_dept"),
            ("责任班组", "team"), ("偏差描述", "description"),
            ("分包单位", "subcontractor"), ("状态", "status"),
        ]),
        ("约谈记录", [
            ("序号", "seq"), ("约谈主题", "name"), ("约谈日期", "record_date"),
            ("被约谈方", "subcontractor"), ("被约谈人", "responsible_person"),
            ("约谈纪要", "description"), ("状态", "status"),
        ]),
    ]

    db_fields = {"name", "record_date", "description", "issue_dept", "responsible_dept", "responsible_person", "area", "subcontractor", "team", "status"}

    for sheet_name, columns in sheet_specs:
        cat_rows = [r for r in rows if r["category"] == sheet_name]
        ws = wb.create_sheet(title=sheet_name)
        ws.sheet_view.showGridLines = False

        # Write header row
        for ci, (header, _) in enumerate(columns, 1):
            cell = ws.cell(row=1, column=ci, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center", vertical="center")
        ws.row_dimensions[1].height = 28

        # Write data rows
        for ri, event in enumerate(cat_rows):
            row_num = ri + 2
            for ci, (_, fname) in enumerate(columns, 1):
                cell = ws.cell(row=row_num, column=ci)
                if fname == "seq":
                    cell.value = ri + 1
                elif fname in db_fields:
                    val = event[fname] or "" if event[fname] else ""
                    if fname == "record_date" and val:
                        try:
                            cell.value = date.fromisoformat(val)
                        except (ValueError, TypeError):
                            cell.value = val
                    else:
                        cell.value = val
                else:
                    cell.value = ""
                cell.font = cell_font
                cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
                cell.border = Border(bottom=thin)
                if ri % 2 == 0:
                    cell.fill = light_fill
            ws.row_dimensions[row_num].height = 30

        # Column widths
        for ci in range(1, len(columns) + 1):
            ws.column_dimensions[get_column_letter(ci)].width = 16

        ws.freeze_panes = "A2"

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    filename = f"外部刹车预警台账_{datetime.now():%Y%m%d}.xlsx"
    return send_file(
        output, as_attachment=True, download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


# ── Team Authorization Dashboard ──────────────────────────────────────

AUTH_FILE_GLOB = "授权列表文件_*.xlsx"


def _parse_auth_counts(filepath):
    """Parse an authorization Excel file, return {team: {auth_name: count}}."""
    wb = load_workbook(filepath, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(min_row=2, values_only=True))
    wb.close()

    # team → auth_name → set of person_ids (for dedup)
    team_auth_persons = {}
    for row in rows:
        name_raw = str(row[0]).strip() if row[0] else ""
        # Extract person ID from "[C517863]裴悦情" format
        pid = name_raw
        team = str(row[5]).strip() if row[5] else ""  # col 6: 班组
        auth = str(row[11]).strip() if row[11] else ""  # col 12: 授权名称
        if not team or not auth:
            continue
        team_auth_persons.setdefault(team, {}).setdefault(auth, set()).add(pid)

    result = {}
    for team, auth_map in team_auth_persons.items():
        result[team] = {auth: len(persons) for auth, persons in auth_map.items()}
    return result


@app.get("/team-auth")
def team_auth_page():
    return send_from_directory("public", "team-auth.html")


@app.post("/api/team-auth/import")
def team_auth_import():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    file = request.files.get("file")
    if not file:
        return jsonify(error="请选择授权列表文件"), 400
    if not file.filename:
        return jsonify(error="文件名无效"), 400
    try:
        filepath = ROOT / file.filename
        file.save(str(filepath))
        counts = _parse_auth_counts(filepath)

        # Build auth type list: priority items first, then sorted by total count desc
        PRIORITY_AUTHS = ["司索", "动火作业监护", "有限空间作业", "倒车指挥", "有限空间作业监护", "手拉葫芦", "起重作业监护", "高处作业监护"]
        auth_totals = {}
        for team, auths in counts.items():
            for auth, cnt in auths.items():
                auth_totals[auth] = auth_totals.get(auth, 0) + cnt

        def _auth_sort_key(auth):
            try:
                return (0, PRIORITY_AUTHS.index(auth))
            except ValueError:
                return (1, -auth_totals.get(auth, 0))

        auth_columns = sorted(auth_totals.keys(), key=_auth_sort_key)

        # Determine thresholds
        def is_over(auth_name, team_name, count):
            is_mao = "铆工" in team_name
            is_gang = "钢筋" in team_name
            is_special_auth = any(k in auth_name for k in ["司索", "角磨机", "千斤顶"])
            if (is_mao or is_gang) and is_special_auth:
                return count > 25
            return count > 15

        # Build matrix
        teams_sorted = sorted(counts.keys())
        matrix = []
        for team in teams_sorted:
            row_data = []
            for auth in auth_columns:
                cnt = counts[team].get(auth, 0)
                over = is_over(auth, team, cnt)
                row_data.append({"auth": auth, "count": cnt, "over": over, "team": team})
            matrix.append({"team": team, "cells": row_data})

        result = {
            "auth_columns": auth_columns,
            "matrix": matrix,
            "auth_totals": auth_totals,
        }

        now = datetime.now().isoformat(timespec="seconds")
        with db() as conn:
            conn.execute("DELETE FROM team_auth_dashboard")
            conn.execute(
                "INSERT INTO team_auth_dashboard (id, auth_filename, result_json, imported_at) VALUES (1,?,?,?)",
                (file.filename, json.dumps(result, ensure_ascii=False), now),
            )

        return jsonify(ok=True, filename=file.filename, **result)
    except Exception as e:
        return jsonify(error=str(e)), 500


@app.get("/api/team-auth/current")
def team_auth_current():
    with db() as conn:
        row = conn.execute("SELECT auth_filename, result_json, imported_at FROM team_auth_dashboard WHERE id=1").fetchone()
    if not row or not row["result_json"]:
        return jsonify(error="尚未导入授权列表数据"), 404
    data = json.loads(row["result_json"])
    data["auth_filename"] = row["auth_filename"]
    data["imported_at"] = row["imported_at"]
    return jsonify(data)


@app.get("/api/team-auth/export-over")
def team_auth_export_over():
    with db() as conn:
        row = conn.execute("SELECT auth_filename, result_json FROM team_auth_dashboard WHERE id=1").fetchone()
    if not row or not row["result_json"]:
        return jsonify(error="暂无数据"), 404
    data = json.loads(row["result_json"])
    matrix = data.get("matrix", [])

    def is_over(auth_name, team_name, count):
        is_mao = "铆工" in team_name
        is_gang = "钢筋" in team_name
        is_special = any(k in auth_name for k in ["司索", "角磨机", "千斤顶"])
        if (is_mao or is_gang) and is_special:
            return count > 25
        return count > 15

    over_rows = []
    for team_row in matrix:
        for cell in team_row["cells"]:
            if cell["count"] > 0 and cell["over"]:
                threshold = 25 if (("铆工" in cell["team"] or "钢筋" in cell["team"]) and any(k in cell["auth"] for k in ["司索", "角磨机", "千斤顶"])) else 15
                over_rows.append({
                    "team": cell["team"],
                    "auth": cell["auth"],
                    "count": cell["count"],
                    "threshold": threshold,
                    "exceed": cell["count"] - threshold,
                })

    over_rows.sort(key=lambda r: (-r["count"], r["team"], r["auth"]))

    wb = Workbook()
    ws = wb.active
    ws.title = "超标条目"
    ws.sheet_view.showGridLines = False

    hdr_font = Font(name="微软雅黑", size=10, bold=True, color="FFFFFF")
    hdr_fill = PatternFill("solid", fgColor="B83B32")
    cell_font = Font(name="微软雅黑", size=10, color="24332E")
    row_fill = PatternFill("solid", fgColor="FFF0EF")
    thin = Side(style="thin", color="E8D0CE")
    center = Alignment(horizontal="center", vertical="center")

    headers = ["班组", "授权名称", "持证人数", "阈值", "超出人数"]
    for ci, h in enumerate(headers, 1):
        c = ws.cell(row=1, column=ci, value=h)
        c.font = hdr_font; c.fill = hdr_fill; c.alignment = center
    ws.row_dimensions[1].height = 26

    for ri, item in enumerate(over_rows):
        rn = ri + 2
        vals = [item["team"], item["auth"], item["count"], item["threshold"], item["exceed"]]
        for ci, v in enumerate(vals, 1):
            c = ws.cell(row=rn, column=ci, value=v)
            c.font = cell_font; c.alignment = center
            c.border = Border(bottom=thin)
            if ri % 2 == 0:
                c.fill = row_fill
        ws.row_dimensions[rn].height = 24

    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 14
    ws.column_dimensions["D"].width = 10
    ws.column_dimensions["E"].width = 14
    ws.freeze_panes = "A2"

    buf = BytesIO()
    wb.save(buf); buf.seek(0)
    now = datetime.now().strftime("%Y%m%d")
    return send_file(buf, as_attachment=True, download_name=f"班组授权超标条目_{now}.xlsx",
                     mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ── Experience Feedback Ledger ────────────────────────────────────────

def _feedback_file_dict(row):
    item = dict(row)
    item["download_url"] = f"/api/feedback/files/{item['id']}/download"
    item["preview_url"] = f"/api/feedback/files/{item['id']}/preview"
    return item


@app.get("/feedback-ledger")
def feedback_page():
    return send_from_directory("public", "feedback-ledger.html")


@app.get("/api/feedback/categories")
def feedback_categories():
    with db() as conn:
        rows = conn.execute("SELECT id, name, sort_order FROM feedback_categories ORDER BY sort_order, id").fetchall()
    return jsonify(items=[dict(row) for row in rows])


@app.post("/api/feedback/categories")
def feedback_create_category():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        if conn.execute("SELECT id FROM feedback_categories WHERE name=?", (name,)).fetchone():
            return jsonify(error="模块名称已存在"), 409
        max_order = conn.execute("SELECT COALESCE(MAX(sort_order), -1) AS o FROM feedback_categories").fetchone()["o"]
        cursor = conn.execute("INSERT INTO feedback_categories (name, sort_order, created_at) VALUES (?, ?, ?)", (name, max_order + 1, now))
    return jsonify(ok=True, id=cursor.lastrowid, name=name, sort_order=max_order + 1)


@app.patch("/api/feedback/categories/<int:category_id>")
def feedback_update_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    if not name:
        return jsonify(error="请输入模块名称"), 400
    with db() as conn:
        cat = conn.execute("SELECT id, name FROM feedback_categories WHERE id=?", (category_id,)).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        if conn.execute("SELECT id FROM feedback_categories WHERE name=? AND id!=?", (name, category_id)).fetchone():
            return jsonify(error="模块名称已存在"), 409
        conn.execute("UPDATE feedback_categories SET name=? WHERE id=?", (name, category_id))
        conn.execute("UPDATE feedback_events SET category=? WHERE category=?", (name, cat["name"]))
    return jsonify(ok=True, id=category_id, name=name)


@app.delete("/api/feedback/categories/<int:category_id>")
def feedback_delete_category(category_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        cat = conn.execute("SELECT id, name FROM feedback_categories WHERE id=?", (category_id,)).fetchone()
        if not cat:
            return jsonify(error="模块不存在"), 404
        if conn.execute("SELECT COUNT(*) AS cnt FROM feedback_categories").fetchone()["cnt"] <= 1:
            return jsonify(error="至少保留一个模块"), 400
        conn.execute("UPDATE feedback_events SET category='经验反馈' WHERE category=?", (cat["name"],))
        conn.execute("DELETE FROM feedback_categories WHERE id=?", (category_id,))
    return jsonify(ok=True, id=category_id)


@app.post("/api/feedback/categories/reorder")
def feedback_reorder_categories():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    ids = body.get("ids")
    if not ids or not isinstance(ids, list):
        return jsonify(error="请提供模块ID列表"), 400
    with db() as conn:
        for index, cid in enumerate(ids):
            conn.execute("UPDATE feedback_categories SET sort_order=? WHERE id=?", (index, int(cid)))
    return jsonify(ok=True)


@app.get("/api/feedback/events")
def feedback_events():
    keyword = (request.args.get("keyword") or "").strip()
    category = (request.args.get("category") or "").strip()
    where = ""
    params = []
    if category:
        where = "WHERE e.category = ?"
        params.append(category)
    if keyword:
        terms = [t for t in re.split(r"\s+", keyword) if t]
        searchable = """(e.name LIKE ? OR e.content LIKE ? OR e.record_date LIKE ?
             OR EXISTS (SELECT 1 FROM feedback_files sf WHERE sf.event_id=e.id AND (sf.original_name LIKE ? OR sf.display_name LIKE ?)))"""
        prefix = "AND " if where else "WHERE "
        where += prefix + " AND ".join(searchable for _ in terms)
        for term in terms:
            params.extend([f"%{term}%"] * 5)
    with db() as conn:
        events = conn.execute(f"""
            SELECT e.*, COUNT(f.id) AS file_count
            FROM feedback_events e
            LEFT JOIN feedback_files f ON f.event_id=e.id
            {where}
            GROUP BY e.id
            ORDER BY e.record_date DESC, e.id DESC
        """, params).fetchall()
        result = []
        for event in events:
            files = conn.execute(
                "SELECT id,event_id,original_name,display_name,kind,content_type,size,created_at FROM feedback_files WHERE event_id=? ORDER BY id DESC",
                (event["id"],),
            ).fetchall()
            entry = dict(event)
            entry["files"] = [_feedback_file_dict(row) for row in files]
            result.append(entry)
    return jsonify(items=result)


@app.post("/api/feedback/events")
def feedback_create_event():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    name = str(body.get("name") or "").strip()
    record_date = str(body.get("record_date") or "").strip()
    category = str(body.get("category") or "经验反馈").strip()
    content = str(body.get("content") or "").strip()
    participant_count = int(body.get("participant_count") or 0)
    if not name:
        return jsonify(error="请输入事件名称"), 400
    try:
        date.fromisoformat(record_date)
    except ValueError:
        return jsonify(error="请选择正确的日期"), 400
    now = datetime.now().isoformat(timespec="seconds")
    with db() as conn:
        cursor = conn.execute(
            "INSERT INTO feedback_events (name,record_date,category,content,participant_count,created_at) VALUES (?,?,?,?,?,?)",
            (name, record_date, category, content, participant_count, now),
        )
    return jsonify(ok=True, id=cursor.lastrowid)


@app.patch("/api/feedback/events/<int:event_id>")
def feedback_update_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    body = request.get_json(force=True)
    fields = {}
    for key in ("name", "record_date", "category", "content"):
        if key in body:
            fields[key] = str(body.get(key) or "").strip()
    if "participant_count" in body:
        fields["participant_count"] = int(body.get("participant_count") or 0)
    if "name" in fields and not fields["name"]:
        return jsonify(error="请输入事件名称"), 400
    if "record_date" in fields:
        try:
            date.fromisoformat(fields["record_date"])
        except ValueError:
            return jsonify(error="请选择正确的日期"), 400
    if not fields:
        return jsonify(ok=True, id=event_id)
    assignments = ",".join(f"{k}=?" for k in fields)
    with db() as conn:
        cursor = conn.execute(f"UPDATE feedback_events SET {assignments} WHERE id=?", [*fields.values(), event_id])
        event = conn.execute("SELECT name,record_date FROM feedback_events WHERE id=?", (event_id,)).fetchone()
        if event:
            files = conn.execute("SELECT id,original_name,kind FROM feedback_files WHERE event_id=?", (event_id,)).fetchall()
            year, month, day = map(int, event["record_date"].split("-"))
            for item in files:
                suffix = Path(item["original_name"]).suffix.lower()
                file_label = "照片" if item["kind"] == "image" else "文件"
                display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
                conn.execute("UPDATE feedback_files SET display_name=? WHERE id=?", (display_name, item["id"]))
    if not cursor.rowcount:
        return jsonify(error="记录不存在"), 404
    return jsonify(ok=True, id=event_id)


@app.delete("/api/feedback/events/<int:event_id>")
def feedback_delete_event(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        event = conn.execute("SELECT id FROM feedback_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="记录不存在"), 404
        stored_names = [row["stored_name"] for row in conn.execute("SELECT stored_name FROM feedback_files WHERE event_id=?", (event_id,))]
        conn.execute("DELETE FROM feedback_events WHERE id=?", (event_id,))
    for name in stored_names:
        target = FEEDBACK_DIR / name
        if target.parent == FEEDBACK_DIR and target.is_file():
            target.unlink()
    return jsonify(ok=True, id=event_id)


@app.post("/api/feedback/events/<int:event_id>/files")
def feedback_upload(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    files = [f for f in request.files.getlist("files") if f and f.filename]
    if not files:
        return jsonify(error="请选择要上传的文件"), 400
    with db() as conn:
        event = conn.execute("SELECT * FROM feedback_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="记录不存在"), 404
        saved = []
        for upload_file in files:
            original = Path(upload_file.filename).name
            safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original).strip(" .") or "文件"
            suffix = Path(safe_original).suffix.lower()
            stored_name = f"{event_id}_{uuid.uuid4().hex}{suffix}"
            target = FEEDBACK_DIR / stored_name
            upload_file.save(target)
            ct = upload_file.mimetype or "application/octet-stream"
            if ct.startswith("image/"):
                kind = "image"
            elif ct == "application/pdf":
                kind = "pdf"
            elif ct in ("application/vnd.openxmlformats-officedocument.wordprocessingml.document",):
                kind = "docx"
            else:
                kind = "file"
            year, month, day = map(int, event["record_date"].split("-"))
            file_label = "照片" if kind == "image" else "文件"
            display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
            now = datetime.now().isoformat(timespec="seconds")
            cursor = conn.execute(
                "INSERT INTO feedback_files (event_id,original_name,stored_name,display_name,kind,content_type,size,created_at) VALUES (?,?,?,?,?,?,?,?)",
                (event_id, original, stored_name, display_name, kind, ct, upload_file.tell(), now),
            )
            saved.append({"id": cursor.lastrowid, "display_name": display_name})
    return jsonify(ok=True, items=saved)


@app.get("/api/feedback/files/<int:file_id>/download")
def feedback_download(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM feedback_files WHERE id=?", (file_id,)).fetchone()
    if not item:
        return jsonify(error="文件不存在"), 404
    target = FEEDBACK_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    return send_file(target, as_attachment=True, download_name=item["display_name"], mimetype=item["content_type"])


@app.get("/api/feedback/files/<int:file_id>/preview")
def feedback_preview(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM feedback_files WHERE id=?", (file_id,)).fetchone()
    if not item:
        return jsonify(error="文件不存在"), 404
    target = FEEDBACK_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    if item["kind"] == "image":
        return send_file(target, mimetype=item["content_type"], conditional=True)
    if item["kind"] == "pdf":
        return send_file(target, mimetype="application/pdf", conditional=True)
    if item["kind"] == "docx":
        try:
            from docx import Document
            doc = Document(target)
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return jsonify(text=text, kind="docx")
        except Exception as e:
            return jsonify(error=f"DOCX 解析失败: {str(e)}"), 500
    return jsonify(error="不支持预览此文件类型"), 400


@app.delete("/api/feedback/files/<int:file_id>")
def feedback_delete_file(file_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        item = conn.execute("SELECT * FROM feedback_files WHERE id=?", (file_id,)).fetchone()
        if not item:
            return jsonify(error="文件不存在"), 404
        conn.execute("DELETE FROM feedback_files WHERE id=?", (file_id,))
    target = FEEDBACK_DIR / item["stored_name"]
    if target.is_file():
        target.unlink()
    return jsonify(ok=True, id=file_id)


@app.get("/api/feedback/stats")
def feedback_stats():
    start_date = (request.args.get("start_date") or "").strip()
    end_date = (request.args.get("end_date") or "").strip()
    category = (request.args.get("category") or "").strip()
    where = []
    params = []
    if start_date:
        where.append("record_date >= ?")
        params.append(start_date)
    if end_date:
        where.append("record_date <= ?")
        params.append(end_date)
    if category:
        where.append("category = ?")
        params.append(category)
    clause = ("WHERE " + " AND ".join(where)) if where else ""
    with db() as conn:
        row = conn.execute(
            f"SELECT COUNT(*) AS records, COALESCE(SUM(participant_count), 0) AS participants, (SELECT COUNT(*) FROM feedback_files{(' WHERE event_id IN (SELECT id FROM feedback_events '+clause+')') if clause else ''}) AS files FROM feedback_events {clause}",
            params,
        ).fetchone()
    return jsonify(records=row["records"], files=row["files"], participants=row["participants"])


@app.post("/api/feedback/import")
def feedback_import():
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    files = [f for f in request.files.getlist("files") if f and f.filename]
    if not files:
        return jsonify(error="请选择要导入的 Excel 文件"), 400

    imported = 0
    skipped = 0
    items = []
    with db() as conn:
        for upload_file in files:
            try:
                file_bytes = upload_file.read()
                wb = load_workbook(BytesIO(file_bytes), read_only=True, data_only=True)
                ws = wb.active
                rows = list(ws.iter_rows(min_row=1, max_row=min(ws.max_row, 2000), values_only=True))
                wb.close()

                if len(rows) < 2:
                    skipped += 1
                    continue

                name = ""
                record_date = ""
                source_ref = ""
                for row in rows[1:]:
                    if not name and len(row) > 4 and row[4]:
                        name = str(row[4]).strip()
                    if not record_date and len(row) > 1 and row[1]:
                        val = str(row[1]).strip()
                        record_date = val[:10] if len(val) >= 10 else val
                    if not source_ref and len(row) > 3 and row[3]:
                        source_ref = str(row[3]).strip()
                    if name and record_date and source_ref:
                        break

                participant_count = len([r for r in rows[1:] if any(c is not None for c in r)])

                if not name:
                    skipped += 1
                    continue

                existing = conn.execute(
                    "SELECT id FROM feedback_events WHERE source_ref = ? AND source_ref != ''",
                    (source_ref,),
                ).fetchone()

                now = datetime.now().isoformat(timespec="seconds")

                if existing:
                    event_id = existing["id"]
                    skipped += 1
                else:
                    cursor = conn.execute(
                        "INSERT INTO feedback_events (name, record_date, category, content, participant_count, source_ref, created_at) VALUES (?,?,?,?,?,?,?)",
                        (name, record_date, "经验反馈", "", participant_count, source_ref, now),
                    )
                    event_id = cursor.lastrowid
                    imported += 1
                    items.append({"name": name, "participant_count": participant_count})

                # Save the uploaded xlsx as an attachment
                original = Path(upload_file.filename).name
                safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original).strip(" .") or "文件"
                suffix = Path(safe_original).suffix.lower()
                stored_name = f"{event_id}_{uuid.uuid4().hex}{suffix}"
                target = FEEDBACK_DIR / stored_name
                target.write_bytes(file_bytes)
                try:
                    year, month, day = map(int, record_date.split("-"))
                except (ValueError, AttributeError):
                    year, month, day = 2026, 1, 1
                display_name = f"{year}年{month}月{day}日{name}文件{suffix}"
                conn.execute(
                    "INSERT INTO feedback_files (event_id, original_name, stored_name, display_name, kind, content_type, size, created_at) VALUES (?,?,?,?,?,?,?,?)",
                    (event_id, original, stored_name, display_name, "file", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", len(file_bytes), now),
                )
            except Exception:
                skipped += 1

    return jsonify(ok=True, imported=imported, skipped=skipped, items=items)


@app.post("/api/feedback/export")
def feedback_export():
    body = request.get_json(force=True) or {}
    raw_ids = body.get("ids") or []
    try:
        event_ids = [int(v) for v in raw_ids]
    except (TypeError, ValueError):
        return jsonify(error="记录选择无效"), 400
    with db() as conn:
        if event_ids:
            placeholders = ",".join("?" for _ in event_ids)
            rows = conn.execute(
                f"SELECT id, name, record_date, participant_count FROM feedback_events WHERE id IN ({placeholders}) ORDER BY record_date, id",
                event_ids,
            ).fetchall()
        else:
            rows = conn.execute(
                "SELECT id, name, record_date, participant_count FROM feedback_events ORDER BY record_date, id"
            ).fetchall()
    if not rows:
        return jsonify(error="未找到所选记录"), 404

    wb = Workbook()
    ws = wb.active
    ws.title = "经验反馈台账"
    ws.sheet_view.showGridLines = False

    header_font = Font(name="微软雅黑", size=11, bold=True, color="FFFFFF")
    header_fill = PatternFill("solid", fgColor="267B67")
    cell_font = Font(name="微软雅黑", size=11, color="24332E")
    thin = Side(style="thin", color="D9E6E1")
    light_fill = PatternFill("solid", fgColor="F3F8F6")

    headers = ["序号", "反馈日期", "反馈主题", "反馈人数"]
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=ci, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 30

    for ri, event in enumerate(rows):
        row_num = ri + 2
        vals = [ri + 1, event["record_date"], event["name"], event["participant_count"] or 0]
        for ci, val in enumerate(vals, 1):
            cell = ws.cell(row=row_num, column=ci, value=val)
            cell.font = cell_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = Border(bottom=thin)
            if ri % 2 == 0:
                cell.fill = light_fill
        ws.row_dimensions[row_num].height = 28

    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 14
    ws.column_dimensions["C"].width = 55
    ws.column_dimensions["D"].width = 12
    ws.freeze_panes = "A2"

    output = BytesIO()
    wb.save(output)
    output.seek(0)
    filename = f"经验反馈台账_{datetime.now():%Y%m%d}.xlsx"
    return send_file(
        output, as_attachment=True, download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.post("/api/training-ledger/events/<int:event_id>/files")
def training_ledger_upload(event_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    files = [item for item in request.files.getlist("files") if item and item.filename]
    if not files:
        return jsonify(error="请选择要上传的文件"), 400
    with db() as conn:
        event = conn.execute("SELECT * FROM training_ledger_events WHERE id=?", (event_id,)).fetchone()
        if not event:
            return jsonify(error="培训名目不存在"), 404
        saved = []
        for upload_file in files:
            original = Path(upload_file.filename).name
            safe_original = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "_", original).strip(" .") or "文件"
            suffix = Path(safe_original).suffix.lower()
            stored_name = f"{event_id}_{uuid.uuid4().hex}{suffix}"
            target = TRAINING_LEDGER_DIR / stored_name
            upload_file.save(target)
            content_type = upload_file.mimetype or "application/octet-stream"
            kind = "image" if content_type.startswith("image/") else "file"
            year, month, day = map(int, event["training_date"].split("-"))
            file_label = "照片" if kind == "image" else "签到单"
            display_name = f"{year}年{month}月{day}日{event['name']}{file_label}{suffix}"
            now = datetime.now().isoformat(timespec="seconds")
            cursor = conn.execute("""
                INSERT INTO training_ledger_files
                (event_id,original_name,stored_name,display_name,kind,content_type,size,created_at)
                VALUES (?,?,?,?,?,?,?,?)
            """, (event_id, original, stored_name, display_name, kind, content_type, target.stat().st_size, now))
            saved.append({"id": cursor.lastrowid, "display_name": display_name})
    return jsonify(ok=True, items=saved)


@app.get("/api/training-ledger/files/<int:file_id>/download")
def training_ledger_download(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM training_ledger_files WHERE id=?", (file_id,)).fetchone()
    if not item:
        return jsonify(error="文件不存在"), 404
    target = TRAINING_LEDGER_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    return send_file(target, as_attachment=True, download_name=item["display_name"], mimetype=item["content_type"])


@app.get("/api/training-ledger/files/<int:file_id>/preview")
def training_ledger_preview(file_id):
    with db() as conn:
        item = conn.execute("SELECT * FROM training_ledger_files WHERE id=?", (file_id,)).fetchone()
    if not item or item["kind"] != "image":
        return jsonify(error="图片不存在"), 404
    target = TRAINING_LEDGER_DIR / item["stored_name"]
    if not target.is_file():
        return jsonify(error="文件已丢失"), 404
    return send_file(target, mimetype=item["content_type"], conditional=True)


@app.delete("/api/training-ledger/files/<int:file_id>")
def training_ledger_delete_file(file_id):
    if not ledger_password_ok():
        return jsonify(error="管理密码错误"), 403
    with db() as conn:
        item = conn.execute("SELECT * FROM training_ledger_files WHERE id=?", (file_id,)).fetchone()
        if not item:
            return jsonify(error="文件不存在"), 404
        conn.execute("DELETE FROM training_ledger_files WHERE id=?", (file_id,))
    target = TRAINING_LEDGER_DIR / item["stored_name"]
    if target.is_file():
        target.unlink()
    return jsonify(ok=True, id=file_id)


@app.get("/api/training/materials")
def training_materials():
    materials = load_training_materials()
    category = request.args.get("category", "").strip()
    if category:
        materials = [m for m in materials if m.get("category") == category]
    materials.sort(key=lambda m: (MATERIAL_CATEGORIES.index(m.get("category", "其他")) if m.get("category") in MATERIAL_CATEGORIES else 99, m.get("sort", 0)))
    return jsonify({"items": materials, "categories": MATERIAL_CATEGORIES})


@app.post("/api/training/materials/save")
def training_materials_save():
    body = request.get_json(force=True)
    if str(body.get("password") or "") != TRAINING_EDIT_PASSWORD:
        return jsonify(error="密码错误"), 403
    materials = load_training_materials()
    item_id = str(body.get("id") or "").strip()
    now = datetime.now().isoformat(timespec="seconds")
    if item_id:
        found = False
        for m in materials:
            if m["id"] == item_id:
                for key in ("category", "title", "content", "sort"):
                    if key in body and body[key] is not None:
                        m[key] = body[key]
                m["updated_at"] = now
                found = True
                break
        if not found:
            return jsonify(error="未找到该资料"), 404
    else:
        materials.append({
            "id": str(uuid.uuid4())[:8],
            "category": body.get("category", "其他"),
            "title": body.get("title", ""),
            "content": body.get("content", ""),
            "sort": int(body.get("sort", 0)),
            "updated_at": now,
        })
    save_training_materials(materials)
    return jsonify(ok=True)


@app.post("/api/training/materials/delete")
def training_materials_delete():
    body = request.get_json(force=True)
    if str(body.get("password") or "") != TRAINING_EDIT_PASSWORD:
        return jsonify(error="密码错误"), 403
    item_id = str(body.get("id") or "").strip()
    materials = [m for m in load_training_materials() if m["id"] != item_id]
    save_training_materials(materials)
    return jsonify(ok=True)


# ---------------------------------------------------------------------------
# Alert dashboard routes
# ---------------------------------------------------------------------------

EXT_TYPE_ORDER = ["挂牌督办", "管理约谈", "红黄牌", "处理通报", "停工令", "整改单", "违章培训通知单", "监理通知单"]
INT_TYPE_ORDER = ["停工令", "处理通报", "整改单", "违章培训通知单"]


@app.get("/alert")
def alert_page():
    return send_from_directory(app.static_folder + "/alert", "index.html")


@app.post("/api/alert/import")
def alert_import():
    file = request.files.get("file")
    if not file or not file.filename:
        return jsonify(error="请选择 Excel 文件"), 400
    if not file.filename.lower().endswith((".xlsx", ".xls")):
        return jsonify(error="请上传 .xlsx 或 .xls 文件"), 400
    file.save(ALERT_DATA_FILE)
    global _alert_data
    _alert_data = None
    try:
        data = load_alert_data()
        if not data:
            return jsonify(error="文件解析失败，请确认上传的是防城港三期安全管理数据总台账"), 400
        return jsonify(ok=True, updated=True,
                       external_count=len(data.get("external", [])),
                       internal_count=len(data.get("internal", [])))
    except Exception as exc:
        ALERT_DATA_FILE.unlink(missing_ok=True)
        _alert_data = {}
        return jsonify(error=f"文件解析失败: {exc}"), 400


def _filter_monthly(items, start, end):
    """Return items with 'total' aggregated from months overlapping [start, end].
    If start/end are empty, return unchanged (all months)."""
    if not start and not end:
        return items
    try:
        s_month = int(start[5:7]) if start else 1
        e_month = int(end[5:7]) if end else 12
    except (ValueError, IndexError):
        return items
    if s_month < 1: s_month = 1
    if e_month > 12: e_month = 12
    if s_month > e_month:
        s_month, e_month = e_month, s_month

    result = []
    for item in items:
        monthly = item.get("monthly", [])
        total = 0
        for m in range(s_month - 1, e_month):
            if m < len(monthly):
                total += monthly[m]
        result.append({**item, "total": total})
    return result


@app.get("/api/alert/summary")
def alert_summary():
    data = load_alert_data()
    start = request.args.get("start", "").strip()
    end = request.args.get("end", "").strip()
    external = _filter_monthly(data.get("external", []), start, end)
    internal = _filter_monthly(data.get("internal", []), start, end)
    return jsonify({
        "external": external,
        "internal": internal,
        "months": data.get("months", []),
        "start": start,
        "end": end,
    })


@app.get("/api/alert/scores")
def alert_scores():
    data = load_alert_data()
    return jsonify(data.get("scores", {}))


def _aggregate_types(type_data, type_order):
    """Convert {type_name: count} to ordered list of {name, count}."""
    result = []
    for t in type_order:
        result.append({"name": t, "count": type_data.get(t, 0)})
    return result


def _aggregate_category(records, start, end, name_filter, category):
    """Aggregate type counts from detail_records filtered by date range."""
    types = {}
    for rec in records:
        if start and rec["date"] < start:
            continue
        if end and rec["date"] > end:
            continue
        if rec["category"] != category:
            continue
        if name_filter and rec["sub_name"] != name_filter and rec["dept_name"] != name_filter:
            continue
        tn = rec["type_name"]
        types[tn] = types.get(tn, 0) + 1
    total = sum(types.values())
    return types, total


def _build_bar_data(records, start, end, field):
    """Build ranking data from detail_records filtered by date range.
    field: 'sub_name' or 'dept_name'"""
    totals = {}
    for rec in records:
        if start and rec["date"] < start:
            continue
        if end and rec["date"] > end:
            continue
        name = rec.get(field, "")
        if not name:
            continue
        if name not in totals:
            totals[name] = {"external": 0, "internal": 0}
        totals[name][rec["category"]] += 1
    result = []
    for name, counts in totals.items():
        result.append({
            "name": name,
            "external": counts["external"],
            "internal": counts["internal"],
            "total": counts["external"] + counts["internal"],
        })
    result.sort(key=lambda x: -x["total"])
    return result


@app.get("/api/alert/departments")
def alert_departments():
    data = load_alert_data()
    start = request.args.get("start", "").strip()
    end = request.args.get("end", "").strip()
    dept = request.args.get("dept", "").strip()

    records = data.get("detail_records", [])
    ext_types, ext_total = _aggregate_category(records, start, end, dept, "external")
    int_types, int_total = _aggregate_category(records, start, end, dept, "internal")

    return jsonify({
        "names": data.get("departments", {}).get("names", []),
        "external": data.get("departments", {}).get("external", []),
        "internal": data.get("departments", {}).get("internal", []),
        "external_types": _aggregate_types(ext_types, EXT_TYPE_ORDER),
        "internal_types": _aggregate_types(int_types, INT_TYPE_ORDER),
        "bar_data": _build_bar_data(records, start, end, "dept_name"),
        "start": start,
        "end": end,
    })


@app.get("/api/alert/subcontractors")
def alert_subcontractors():
    data = load_alert_data()
    start = request.args.get("start", "").strip()
    end = request.args.get("end", "").strip()
    sub = request.args.get("sub", "").strip()

    records = data.get("detail_records", [])
    ext_types, ext_total = _aggregate_category(records, start, end, sub, "external")
    int_types, int_total = _aggregate_category(records, start, end, sub, "internal")

    bar_data = _build_bar_data(records, start, end, "sub_name")
    items = []
    for b in bar_data:
        items.append({"name": b["name"], "count": b["total"],
                       "external": b["external"], "internal": b["internal"]})

    return jsonify({
        "items": items,
        "external_types": _aggregate_types(ext_types, EXT_TYPE_ORDER),
        "internal_types": _aggregate_types(int_types, INT_TYPE_ORDER),
        "bar_data": bar_data,
        "start": start,
        "end": end,
    })


@app.get("/api/alert/type-stats")
def alert_type_stats():
    """Aggregate problem type statistics for pie/bar charts with optional filters."""
    data = load_alert_data()
    start = request.args.get("start", "").strip()
    end = request.args.get("end", "").strip()
    category = request.args.get("category", "").strip()
    brake_type = request.args.get("type", "").strip()  # 预警刹车类型筛选

    records = data.get("detail_records", [])
    # Filter by date range, category, and brake type
    filtered = []
    for rec in records:
        if start and rec["date"] < start:
            continue
        if end and rec["date"] > end:
            continue
        if category and rec["category"] != category:
            continue
        if brake_type and rec["type_name"] != brake_type:
            continue
        # Only include records that have a problem_type
        if not rec.get("problem_type"):
            continue
        filtered.append(rec)

    # Aggregate by problem_type
    type_counts = {}
    for rec in filtered:
        pt = rec["problem_type"]
        cat = rec["category"]
        key = pt
        if key not in type_counts:
            type_counts[key] = {"name": pt, "count": 0, "external": 0, "internal": 0}
        type_counts[key]["count"] += 1
        if cat == "external":
            type_counts[key]["external"] += 1
        else:
            type_counts[key]["internal"] += 1

    type_list = sorted(type_counts.values(), key=lambda x: -x["count"])

    # Additional breakdowns for drill-down
    dept_counts = {}
    sub_counts = {}
    for rec in filtered:
        if rec["dept_name"]:
            dept_counts[rec["dept_name"]] = dept_counts.get(rec["dept_name"], 0) + 1
        if rec["sub_name"]:
            sub_counts[rec["sub_name"]] = sub_counts.get(rec["sub_name"], 0) + 1
    dept_list = sorted([{"name": k, "count": v} for k, v in dept_counts.items()], key=lambda x: -x["count"])
    sub_list = sorted([{"name": k, "count": v} for k, v in sub_counts.items()], key=lambda x: -x["count"])

    return jsonify({
        "type_counts": type_list,
        "dept_counts": dept_list,
        "sub_counts": sub_list,
        "start": start,
        "end": end,
    })


@app.get("/api/alert/details")
def alert_details():
    data = load_alert_data()
    detail_type = request.args.get("type", "").strip()
    sheet_map = {
        "挂牌督办": "工程公司挂牌督办单",
        "管理约谈": "工程公司管理约谈",
        "红黄牌": "红黄牌",
        "处理通报": "工程公司处理通报",
        "通报批评": "工程公司通报批评",
        "工程整改单": "工程公司整改单",
        "工程停工令": "工程公司停工令",
        "监理通知单": "监理业主整改通知单",
        "项目处理通报": "项目内部处理通报",
        "项目整改单": "项目整改通知单",
        "项目停工令": "项目停工令",
        "项目违章培训": "项目违章培训通知单",
        "工程违章培训": "工程公司违章培训通知单",
    }
    sheet_name = sheet_map.get(detail_type)
    if not sheet_name:
        return jsonify({"items": [], "total": 0})

    src = ALERT_DATA_FILE if ALERT_DATA_FILE.exists() else ALERT_SEED_FILE
    if not src.exists():
        return jsonify({"items": [], "total": 0})
    wb = load_workbook(src, data_only=True)
    if sheet_name not in wb.sheetnames:
        wb.close()
        return jsonify({"items": [], "total": 0})

    ws = wb[sheet_name]
    headers = []
    for c in range(1, ws.max_column + 1):
        h = ws.cell(row=1, column=c).value
        if h:
            headers.append(str(h).strip())

    items = []
    for r in range(2, ws.max_row + 1):
        row_vals = [str(ws.cell(row=r, column=c).value or "") for c in range(1, ws.max_column + 1)]
        if all(v == "" for v in row_vals):
            continue
        entry = {}
        for i, h in enumerate(headers):
            if i < len(row_vals):
                entry[h] = row_vals[i]
        items.append(entry)

    wb.close()
    page = int(request.args.get("page", 1))
    size = int(request.args.get("size", 20))
    total = len(items)
    start = (page - 1) * size
    return jsonify({"items": items[start:start + size], "total": total, "page": page, "size": size,
                    "headers": headers})


app.register_blueprint(admin_bp)
app.register_blueprint(meeting_bp)
app.register_blueprint(collection_bp)

init_db()
seed_people()
apply_role_rules()
seed_data()
load_alert_data()

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("APP_PORT", "8010")), debug=os.environ.get("FLASK_DEBUG") == "1")
